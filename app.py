import streamlit as st
from pathlib import Path
import re
import logging
from io import BytesIO
import zipfile
import pytesseract
from PIL import Image
import io
import os
import requests
import gspread
from datetime import datetime
import json
from docx import Document
from typing import List, Dict, Tuple, Optional, Union, Any, Callable
import time
import threading
import concurrent.futures
import google.generativeai as genai

# Set up Gemini API Key
genai.configure(api_key=st.secrets["gemini"]["api_key"])

# def analyze_text_with_gemini(prompt):
#     try:
#         model = genai.GenerativeModel("gemini-2.0-flash")  # Use the appropriate Gemini model
#         response = model.generate_content(prompt)
#         return response.text.strip() if response.text else "No response generated."
#     except Exception as e:
#         return f"Error analyzing text: {str(e)}"

# # Configure logging
# logging.basicConfig(
#     level=logging.INFO,
#     format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
# )
# logger = logging.getLogger(__name__)

# # Import utilities with error handling
# try:
#     from utils.pdf_processing import extract_text_from_pdf
# except ImportError as e:
#     logger.error(f"Failed to import required modules: {str(e)}")
#     st.error(f"Failed to import required modules: {str(e)}")
    
# # Constants
# SPREADSHEET_ID = "1EQiw3jPnTamigx9ZHqOKr6zLl9FlQmko7IQnjlfRtqo"
# CREDENTIALS_PATH = "credentials.json"
# HACKMD_API_TOKEN = "1PQ8Z9IMDPD579ETDEMLJYG8W3V2CY6IPLKYVD7Z8MZQW9WPX9"  # Consider moving to secrets
# MAX_WORKERS = 4  # For parallel processing
# MAX_RETRIES = 3  # For API calls
# SUPPORTED_FILE_TYPES = ["pdf", "docx", "zip", "jpg", "jpeg", "png", "txt"]
# CACHE_TTL = 3600  # Cache time-to-live in seconds
# HACKMD_MAX_CONTENT_LENGTH = 90000  # Reduced size for safer margin (80KB)


# class RateLimitHandler:
#     """Handles rate limiting for API calls with automatic backoff and retry"""
    
#     def __init__(self, 
#                 requests_per_minute: int = 15,
#                 retry_delay: int = 65,  # 65 seconds to ensure full minute has passed
#                 max_retries: int = 5):
#         self.requests_per_minute = requests_per_minute
#         self.retry_delay = retry_delay
#         self.max_retries = max_retries
#         self.request_timestamps = []
        
#     def _clean_old_timestamps(self):
#         """Remove timestamps older than 1 minute"""
#         current_time = time.time()
#         self.request_timestamps = [ts for ts in self.request_timestamps 
#                                if current_time - ts < 60]
        
#     def _can_make_request(self) -> bool:
#         """Check if we can make a request based on the rate limit"""
#         self._clean_old_timestamps()
#         return len(self.request_timestamps) < self.requests_per_minute
        
#     def _time_until_next_available_slot(self) -> float:
#         """Calculate time until we can make another request"""
#         if self._can_make_request():
#             return 0
            
#         self._clean_old_timestamps()
#         oldest_timestamp = min(self.request_timestamps)
#         return max(0, 60 - (time.time() - oldest_timestamp))
        
#     def execute_with_rate_limit(self, 
#                               func: Callable, 
#                               *args, 
#                               **kwargs) -> Any:
#         """
#         Execute a function with rate limiting
#         Returns the function result or raises the last exception after max retries
#         """
#         retries = 0
#         last_exception = None
        
#         while retries <= self.max_retries:
#             try:
#                 # Check if we can make a request
#                 wait_time = self._time_until_next_available_slot()
                
#                 if wait_time > 0:
#                     # Need to wait before making another request
#                     retries += 1
#                     if retries > self.max_retries:
#                         break
                        
#                     st.warning(f"Rate limit reached. Waiting {int(wait_time)} seconds before continuing...")
#                     time.sleep(wait_time)
#                     continue
                
#                 # Record this request timestamp
#                 self.request_timestamps.append(time.time())
                
#                 # Execute the function
#                 return func(*args, **kwargs)
                
#             except Exception as e:
#                 last_exception = e
#                 error_message = str(e).lower()
                
#                 # Check if it's a rate limit error
#                 if "quota" in error_message or "rate" in error_message or "limit" in error_message:
#                     retries += 1
#                     if retries > self.max_retries:
#                         break
                        
#                     st.warning(f"API rate limit exceeded. Waiting {self.retry_delay} seconds before retry {retries}/{self.max_retries}...")
#                     time.sleep(self.retry_delay)
#                 else:
#                     # If it's not a rate limit error, re-raise immediately
#                     raise
        
#         # If we get here, we've exceeded max retries
#         error_msg = f"Failed after {self.max_retries} retries: {str(last_exception)}"
#         st.error(error_msg)
#         raise Exception(error_msg)


# class BatchConceptsProcessor:
#     """Processes concepts in batches to manage API rate limits"""
    
#     def __init__(self, 
#                 batch_size: int = 10, 
#                 batch_delay: int = 65,
#                 api_function: Callable = None,
#                 progress_callback: Callable = None):
#         self.batch_size = batch_size
#         self.batch_delay = batch_delay
#         self.api_function = api_function
#         self.progress_callback = progress_callback
#         self.results = {}
#         self.errors = {}
        
#     def process_concepts(self, 
#                        concepts: List[str], 
#                        prompt_template: str) -> Dict[str, str]:
#         """Process concepts in batches with rate limit handling"""
        
#         # Create batches of concepts
#         batches = []
#         current_batch = []
        
#         for concept in concepts:
#             current_batch.append(concept)
#             if len(current_batch) >= self.batch_size:
#                 batches.append(current_batch)
#                 current_batch = []
                
#         # Add the last batch if it's not empty
#         if current_batch:
#             batches.append(current_batch)
            
#         # Process each batch
#         total_processed = 0
        
#         # Create a status container
#         status_container = st.container()
        
#         for batch_idx, batch in enumerate(batches):
#             with status_container:
#                 st.info(f"Processing batch {batch_idx+1}/{len(batches)} ({len(batch)} concepts)")
                
#                 batch_progress = st.progress(0)
                
#                 # Process each concept in the batch
#                 for i, concept in enumerate(batch):
#                     # Format the prompt for this concept
#                     prompt = prompt_template.format(concept=concept)
                    
#                     try:
#                         # Call the API function
#                         result = self.api_function(prompt)
#                         self.results[concept] = result
#                     except Exception as e:
#                         error_msg = str(e)
#                         logger.error(f"Error processing concept '{concept}': {error_msg}")
#                         self.errors[concept] = error_msg
#                         st.warning(f"Error processing concept: {concept}")
                        
#                     # Update progress
#                     total_processed += 1
#                     batch_progress.progress((i + 1) / len(batch))
                    
#                     if self.progress_callback:
#                         self.progress_callback(total_processed, len(concepts))
                
#                 # Reset the batch progress
#                 batch_progress.empty()
                
#             # Wait between batches to avoid rate limits
#             if batch_idx < len(batches) - 1:  # Don't wait after the last batch
#                 with status_container:
#                     wait_text = st.text(f"Waiting {self.batch_delay} seconds for rate limit to reset...")
#                     wait_progress = st.progress(0)
                    
#                     # Show a countdown timer
#                     for second in range(self.batch_delay):
#                         wait_progress.progress((second + 1) / self.batch_delay)
#                         wait_text.text(f"Waiting for rate limit to reset... {self.batch_delay - second} seconds remaining")
#                         time.sleep(1)
                    
#                     wait_text.empty()
#                     wait_progress.empty()
        
#         # Return all processed results
#         return self.results


# def solution_prompt_for_concept(concept: str) -> str:
#     """Generate a solution prompt for a concept"""
#     return f"""
#     Provide a theoretical answer for the concept: {concept}, ensuring key details are covered without including a summary or conclusion.

#     **For Coding-Related Questions**:
#     - Do NOT provide complete code.
#     - Instead, describe the approach in a structured manner:
#     - Explain the problem briefly.
#     - Outline the key steps needed to solve it.
#     - Mention important considerations like time complexity, edge cases, and best practices.

#     **For Non-Coding Theoretical Concepts**:
#     - Provide an in-depth explanation covering relevant details.
#     - Use bullet points or structured formatting where appropriate.

#     Concept: {concept}
#     """


# class GoogleSheetLogger:
#     """Class to handle Google Sheets logging operations"""
    
#     def __init__(self, spreadsheet_id: str):
#         self.spreadsheet_id = spreadsheet_id
#         self._gc = None
        
#     @property
#     def gc(self):
#         """Lazy loading of Google Sheets client"""
#         if self._gc is None:
#             try:
#                 # Load credentials from Streamlit secrets
#                 credentials = st.secrets["google_service_account"]
#                 self._gc = gspread.service_account_from_dict(credentials)
#             except Exception as e:
#                 logger.error(f"Failed to initialize Google Sheets client: {str(e)}")
#                 st.error(f"Failed to initialize Google Sheets client: {str(e)}")
#                 raise
#         return self._gc
    
#     def log_document(self, 
#                     job_type: str, 
#                     created_at: str, 
#                     prepared_by: str, 
#                     job_id: str, 
#                     document_type: str, 
#                     title_name: str, 
#                     hackmd_link: str, 
#                     interview_round: Optional[str] = None) -> bool:
#         """Log document information to Google Sheets"""
#         try:
#             # Open the spreadsheet and select the "Reference Docs" sheet
#             sheet = self.gc.open_by_key(self.spreadsheet_id).worksheet("Reference Docs")

#             # Prepare the row to include all fields, adding interview_round if present
#             row_data = [job_type, created_at, prepared_by, job_id, document_type, title_name, hackmd_link]
#             if interview_round:
#                 row_data.append(interview_round)

#             # Append the row to the specific sheet
#             sheet.append_row(row_data)
#             logger.info(f"Successfully logged document: {title_name}")
#             return True
#         except Exception as e:
#             logger.error(f"Failed to log data to Google Sheets: {str(e)}")
#             st.error(f"Failed to log data to Google Sheets: {str(e)}")
#             return False


# class HackMDUploader:
#     """Class to handle HackMD upload operations with content splitting"""
    
#     def __init__(self, api_token: str):
#         self.api_token = api_token
#         self.headers = {
#             "Authorization": f"Bearer {self.api_token}",
#             "Content-Type": "application/json"
#         }
#         self.max_content_length = HACKMD_MAX_CONTENT_LENGTH
        
#     @staticmethod
#     def _add_retry_mechanism(func):
#         """Decorator to add retry mechanism to API calls"""
#         def wrapper(*args, **kwargs):
#             for attempt in range(MAX_RETRIES):
#                 try:
#                     return func(*args, **kwargs)
#                 except requests.RequestException as e:
#                     if attempt < MAX_RETRIES - 1:
#                         wait_time = 2 ** attempt  # Exponential backoff
#                         logger.warning(f"Attempt {attempt+1} failed: {str(e)}. Retrying in {wait_time}s...")
#                         time.sleep(wait_time)
#                     else:
#                         logger.error(f"All {MAX_RETRIES} attempts failed: {str(e)}")
#                         raise
#         return wrapper
    
#     @_add_retry_mechanism
#     def _create_single_note(self, content: str, title: str) -> Optional[str]:
#         """Create a single HackMD note and return its URL"""
#         try:
#             create_url = "https://api.hackmd.io/v1/notes"
            
#             data = {
#                 "content": content,
#                 "title": title,
#                 "readPermission": "guest",
#                 "writePermission": "owner",
#                 "publishType": "freely",
#                 "commentPermission": "disabled"
#             }

#             response = requests.post(create_url, json=data, headers=self.headers)

#             if response.status_code == 201:
#                 note_data = response.json()
#                 publish_link = note_data.get("publishLink", None)
                
#                 if publish_link:
#                     logger.info(f"Successfully uploaded to HackMD: {title}")
#                     return publish_link
#                 else:
#                     logger.error("Publish link not found in the response.")
#                     return None
#             else:
#                 logger.error(f"Failed to create note. Status Code: {response.status_code}, Response: {response.text}")
#                 return None
                
#         except Exception as e:
#             logger.error(f"An error occurred while uploading to HackMD: {str(e)}")
#             return None
    
#     def _split_content(self, content: str) -> List[str]:
#         """Split content into chunks of acceptable size with improved handling for large sections"""
#         # Reduce the maximum chunk size to be safer
#         safe_max_size = int(self.max_content_length * 0.8)  # 80% of limit as safety margin
        
#         chunks = []
#         current_chunk = ""
        
#         # First try splitting by paragraphs (standard approach)
#         paragraphs = content.split("\n\n")
        
#         for paragraph in paragraphs:
#             # If a single paragraph is too large, we need to split it further
#             if len(paragraph) > safe_max_size:
#                 # If this is a large section, split it by lines
#                 lines = paragraph.split("\n")
                
#                 # Reset the current sub-chunk
#                 sub_chunk = ""
                
#                 for line in lines:
#                     # If adding this line would exceed the limit, start a new chunk
#                     if len(sub_chunk) + len(line) + 2 > safe_max_size:
#                         if sub_chunk:
#                             chunks.append(sub_chunk)
#                         sub_chunk = line + "\n"
#                     else:
#                         sub_chunk += line + "\n"
                
#                 # Add the last sub-chunk if it's not empty
#                 if sub_chunk:
#                     chunks.append(sub_chunk)
#             # Normal case - add paragraph if it fits, otherwise start new chunk
#             elif len(current_chunk) + len(paragraph) + 2 > safe_max_size:
#                 chunks.append(current_chunk)
#                 current_chunk = paragraph + "\n\n"
#             else:
#                 current_chunk += paragraph + "\n\n"
        
#         # Add the last chunk if it's not empty
#         if current_chunk:
#             chunks.append(current_chunk)
        
#         # Final safety check - ensure no chunk exceeds the limit
#         final_chunks = []
#         for chunk in chunks:
#             if len(chunk) > safe_max_size:
#                 # Split by newlines if still too large
#                 lines = chunk.split("\n")
#                 current_subchunk = ""
                
#                 for line in lines:
#                     if len(current_subchunk) + len(line) + 1 > safe_max_size:
#                         final_chunks.append(current_subchunk)
#                         current_subchunk = line + "\n"
#                     else:
#                         current_subchunk += line + "\n"
                        
#                 if current_subchunk:
#                     final_chunks.append(current_subchunk)
#             else:
#                 final_chunks.append(chunk)
                
#         return final_chunks
    
#     @_add_retry_mechanism
#     def _get_note_content(self, note_url: str) -> Optional[str]:
#         """Get the content of a HackMD note by its URL"""
#         try:
#             # Extract note ID from URL
#             note_id = note_url.split("/")[-1]
            
#             # Fetch note content
#             api_url = f"https://api.hackmd.io/v1/notes/{note_id}"
#             response = requests.get(api_url, headers=self.headers)
            
#             if response.status_code == 200:
#                 return response.json().get("content")
#             else:
#                 logger.error(f"Failed to get note content. Status Code: {response.status_code}")
#                 return None
#         except Exception as e:
#             logger.error(f"Error getting note content: {str(e)}")
#             return None
            
#     @_add_retry_mechanism
#     def _update_note_content(self, note_url: str, new_content: str) -> bool:
#         """Update the content of a HackMD note with improved handling for 202 responses"""
#         try:
#             # Extract note ID from URL
#             note_id = note_url.split("/")[-1]
            
#             # Update note content
#             api_url = f"https://api.hackmd.io/v1/notes/{note_id}"
#             data = {"content": new_content}
            
#             response = requests.patch(api_url, json=data, headers=self.headers)
            
#             if response.status_code == 200:
#                 logger.info(f"Successfully updated note: {note_url}")
#                 return True
#             elif response.status_code == 202:
#                 # Status code 202 means "Accepted" - the request is being processed
#                 logger.info(f"Update request accepted, waiting for processing: {note_url}")
                
#                 # Wait for the update to complete
#                 time.sleep(3)  # Wait 3 seconds before considering it successful
#                 return True
#             else:
#                 logger.error(f"Failed to update note. Status Code: {response.status_code}")
#                 return False
#         except Exception as e:
#             logger.error(f"Error updating note content: {str(e)}")
#             return False
    
#     @_add_retry_mechanism
#     def upload(self, content: str, title_name: str) -> Optional[str]:
#         """Upload content to HackMD, splitting if necessary, and return the published link"""
#         try:
#             # Check if content is too long
#             if len(content) > self.max_content_length:
#                 # Split content into multiple parts
#                 chunks = self._split_content(content)
                
#                 # Show a progress bar for splitting
#                 with st.spinner(f"Content is large. Splitting into {len(chunks)} parts..."):
#                     part_links = []
#                     # Create a simpler index with just the document parts
#                     index_content = f"# {title_name}\n\n## Document Parts\n\n"
                    
#                     # Create each part
#                     for i, chunk in enumerate(chunks):
#                         part_title = f"{title_name} - Part {i+1}"
#                         part_content = f"# {part_title}\n\n{chunk}"
                        
#                         # Upload this part
#                         part_link = self._create_single_note(part_content, part_title)
                        
#                         if part_link:
#                             part_links.append((part_title, part_link))
#                             index_content += f"- [{part_title}]({part_link})\n"
#                         else:
#                             # If upload fails, try with a smaller chunk
#                             logger.error(f"Failed to upload part {i+1}")
#                             st.error(f"Failed to upload part {i+1}")
                            
#                             # Try again with a smaller chunk size
#                             if len(chunk) > 10000:  # Only try splitting if the chunk is reasonably large
#                                 st.warning(f"Trying to split part {i+1} into smaller segments...")
                                
#                                 # Split in half
#                                 half_size = len(chunk) // 2
#                                 first_half = chunk[:half_size]
#                                 second_half = chunk[half_size:]
                                
#                                 # Try uploading first half
#                                 first_half_title = f"{title_name} - Part {i+1}a"
#                                 first_half_content = f"# {first_half_title}\n\n{first_half}"
#                                 first_half_link = self._create_single_note(first_half_content, first_half_title)
                                
#                                 if first_half_link:
#                                     part_links.append((first_half_title, first_half_link))
#                                     index_content += f"- [{first_half_title}]({first_half_link})\n"
                                
#                                 # Try uploading second half
#                                 second_half_title = f"{title_name} - Part {i+1}b"
#                                 second_half_content = f"# {second_half_title}\n\n{second_half}"
#                                 second_half_link = self._create_single_note(second_half_content, second_half_title)
                                
#                                 if second_half_link:
#                                     part_links.append((second_half_title, second_half_link))
#                                     index_content += f"- [{second_half_title}]({second_half_link})\n"
                    
#                     # Create index note with links to all parts - REMOVED the Navigation Links section
#                     index_link = self._create_single_note(index_content, title_name)
                    
#                     if index_link:
#                         st.success(f"Content was split into {len(chunks)} parts due to size limits.")
#                         return index_link
#                     else:
#                         # If index creation fails, return the first part link
#                         if part_links:
#                             st.warning("Created multiple parts but failed to create index. Using first part as main link.")
#                             return part_links[0][1]
#                         else:
#                             return None
#             else:
#                 # Content is short enough for a single note
#                 content_with_title = f"# {title_name}\n\n{content}"
#                 return self._create_single_note(content_with_title, title_name)
                
#         except Exception as e:
#             logger.error(f"An error occurred while splitting and uploading to HackMD: {str(e)}")
#             st.error(f"An error occurred while uploading to HackMD: {str(e)}")
#             return None

# class TextExtractor:
#     """Class to handle text extraction from various file types"""
    
#     @staticmethod
#     def _ensure_tesseract():
#         """Check if Tesseract is installed and configured"""
#         try:
#             pytesseract.get_tesseract_version()
#         except Exception as e:
#             logger.error(f"Tesseract not properly configured: {str(e)}")
#             st.error("Tesseract OCR is not properly configured. Images cannot be processed.")
#             return False
#         return True
            
#     @staticmethod
#     def from_pdf(file_obj) -> str:
#         """Extract text from PDF file"""
#         try:
#             return extract_text_from_pdf(file_obj)
#         except Exception as e:
#             logger.error(f"Error extracting text from PDF: {str(e)}")
#             raise ValueError(f"Error extracting text from PDF: {str(e)}")
            
#     @staticmethod
#     def from_docx(file_obj) -> str:
#         """Extract text from DOCX file"""
#         try:
#             doc = Document(file_obj)
#             return '\n'.join([para.text for para in doc.paragraphs])
#         except Exception as e:
#             logger.error(f"Error extracting text from DOCX: {str(e)}")
#             raise ValueError(f"Error extracting text from DOCX: {str(e)}")
            
#     @staticmethod
#     def from_image(file_obj) -> str:
#         """Extract text from image using OCR"""
#         if not TextExtractor._ensure_tesseract():
#             return ""
            
#         try:
#             img = Image.open(file_obj)
#             return pytesseract.image_to_string(img)
#         except Exception as e:
#             logger.error(f"Error extracting text from image: {str(e)}")
#             raise ValueError(f"Error extracting text from image: {str(e)}")
            
#     @staticmethod
#     def from_zip(file_obj) -> str:
#         """Extract text from all supported files in a ZIP archive"""
#         if not TextExtractor._ensure_tesseract():
#             return ""
            
#         text = ""
#         try:
#             with zipfile.ZipFile(file_obj, "r") as zip_ref:
#                 for file_name in zip_ref.namelist():
#                     ext = file_name.split(".")[-1].lower() if "." in file_name else ""
                    
#                     if ext == "txt":
#                         with zip_ref.open(file_name) as file:
#                             text += file.read().decode('utf-8', errors='replace') + "\n"
#                     elif ext in ["jpg", "jpeg", "png"]:
#                         with zip_ref.open(file_name) as file:
#                             img = Image.open(file)
#                             text += pytesseract.image_to_string(img) + "\n"
#                     elif ext == "pdf":
#                         with zip_ref.open(file_name) as file:
#                             # Need to copy to BytesIO since extract_text_from_pdf expects a file-like object
#                             pdf_bytes = BytesIO(file.read())
#                             text += TextExtractor.from_pdf(pdf_bytes) + "\n"
#                     elif ext == "docx":
#                         with zip_ref.open(file_name) as file:
#                             docx_bytes = BytesIO(file.read())
#                             text += TextExtractor.from_docx(docx_bytes) + "\n"
#             return text
#         except Exception as e:
#             logger.error(f"Error extracting text from ZIP: {str(e)}")
#             raise ValueError(f"Error extracting text from ZIP: {str(e)}")
            
#     @staticmethod
#     def from_txt(file_obj) -> str:
#         """Extract text from plain text file"""
#         try:
#             return file_obj.read().decode('utf-8', errors='replace')
#         except Exception as e:
#             logger.error(f"Error extracting text from TXT: {str(e)}")
#             raise ValueError(f"Error extracting text from TXT: {str(e)}")
            
#     @staticmethod
#     def extract(file_obj, file_type: str) -> str:
#         """Extract text from file based on its type"""
#         file_type = file_type.lower()
        
#         if file_type not in SUPPORTED_FILE_TYPES:
#             logger.error(f"Unsupported file type: {file_type}")
#             raise ValueError(f"Unsupported file type: {file_type}")
            
#         extractors = {
#             "pdf": TextExtractor.from_pdf,
#             "docx": TextExtractor.from_docx,
#             "zip": TextExtractor.from_zip,
#             "jpg": TextExtractor.from_image,
#             "jpeg": TextExtractor.from_image,
#             "png": TextExtractor.from_image,
#             "txt": TextExtractor.from_txt
#         }
        
#         return extractors[file_type](file_obj)


# class DocumentGenerator:
#     """Base class for document generators"""
    
#     def __init__(self, 
#                 title_name: str, 
#                 job_id: str, 
#                 prepared_by: str, 
#                 job_type: str,
#                 interview_round: Optional[str] = None):
#         self.title_name = title_name
#         self.job_id = job_id
#         self.prepared_by = prepared_by
#         self.job_type = job_type
#         self.interview_round = interview_round
#         self.gs_logger = GoogleSheetLogger(SPREADSHEET_ID)
#         self.hackmd_uploader = HackMDUploader(HACKMD_API_TOKEN)
        
#     def _validate_parameters(self) -> bool:
#         """Validate input parameters"""
#         if not self.title_name.strip():
#             st.error("Please enter a valid Title Name before proceeding.")
#             return False
            
#         if not self.job_id.strip():
#             st.error("Please enter a valid Job ID before proceeding.")
#             return False
            
#         if not self.prepared_by.strip():
#             st.error("Please enter the name in 'Prepared By' field before proceeding.")
#             return False
            
#         if not self.job_type.strip():
#             st.error("Please select a valid Job Type before proceeding.")
#             return False
            
#         return True
        
#     def generate(self, *args, **kwargs) -> Optional[str]:
#         """Generate document content - to be implemented by subclasses"""
#         raise NotImplementedError("Subclasses must implement this method")
        
#     def upload_and_log(self, content: str, document_type: str) -> Optional[str]:
#         """Upload document to HackMD and log to Google Sheets"""
#         if not content:
#             logger.error("No content to upload")
#             return None
            
#         hackmd_link = self.hackmd_uploader.upload(content, self.title_name)
        
#         if hackmd_link:
#             created_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
#             self.gs_logger.log_document(
#                 self.job_type,
#                 created_at,
#                 self.prepared_by,
#                 self.job_id,
#                 document_type,
#                 self.title_name,
#                 hackmd_link,
#                 self.interview_round
#             )
            
#             # Display link information in a more user-friendly way
#             st.success("Document successfully generated and uploaded to HackMD!")
            
#             st.markdown("#### HackMD Link:")
#             link_col1, link_col2 = st.columns([3, 1])
            
#             with link_col1:
#                 st.text_area("Copy the link below:", hackmd_link, height=70, key="hackmd_link_text")
            
#             with link_col2:
#                 st.markdown(f"[Open in Browser]({hackmd_link})")
            
#             # If this was a split document, provide additional information
#             if "Part" in hackmd_link:
#                 st.info("""
#                 Note: Due to the large size of your document, it was automatically split into multiple parts.
#                 The link above points to an index page that contains links to all parts of your document.
#                 """)
                
#         return hackmd_link


# class AssignmentDocumentGenerator(DocumentGenerator):
#     """Generator for Assignment Reference Documents"""
    
#     def generate(self, uploaded_file) -> Optional[str]:
#         """Generate Assignment Reference Document"""
#         if not self._validate_parameters():
#             return None
            
#         document_type = "Assignment Reference Document"
#         file_type = uploaded_file.name.split(".")[-1].lower()
        
#         # Progress bar
#         progress_bar = st.progress(0)
#         status_text = st.empty()
        
#         try:
#             # Extract text
#             status_text.text('Reading and analyzing file...')
#             progress_bar.progress(10)
            
#             pdf_text = TextExtractor.extract(uploaded_file, file_type)
#             if not pdf_text or pdf_text.strip() == "":
#                 st.error(f"Failed to extract text from {uploaded_file.name}.")
#                 return None
                
#             progress_bar.progress(30)
#             status_text.text('Text extracted successfully!')
            
#             # Analyze text
#             status_text.text('Preparing Reference Document...')
#             prompt = f""" 
#             You are tasked with creating a detailed reference document based on the following assignment description:

#             Assignment Description:
#             {pdf_text}

#             The reference document should include the following sections:

#             Objective: Provide a concise overview of the assignment's purpose and goals. Clearly explain the main objectives and expected learning outcomes.

#             Step-by-Step Instructions: Create detailed, easy-to-follow steps for completing the assignment. Each step should be actionable and include:
#             - Project Setup and Initialization: Describe how to set up the project directory, initialize the project, and install dependencies.
#             - Development Process: Outline detailed steps for Briefly Explained in detail way in lengthy How to develop the Assignment.
#             - Styling and Design: Provide instructions for implementing styling.
#             - Deployment: Describe the process of deploying the project to a suitable platform and configuring the environment.

#             Submission Guidelines: Clearly define how the assignment should be submitted, including the expected format, any required documentation.

#             *Important Notes:*
#             - Avoid using a title like "Reference Document."
#             - Do not include any code snippets.
#             - Do not include a summary or conclusion in the reference document.

#             Generate the content in a structured and professional manner, ensuring clarity and usability.
#             """
            
#             # Create rate limit handler for API call
#             rate_limiter = RateLimitHandler(
#                 requests_per_minute=15,  # Gemini's limit
#                 retry_delay=65,          # Wait a bit over a minute to be safe
#                 max_retries=5            # Allow up to 5 retries
#             )
            
#             # Use rate limiter for API call
#             analyzed_text = rate_limiter.execute_with_rate_limit(analyze_text_with_gemini, prompt)
            
#             if not analyzed_text or analyzed_text.strip() == "":
#                 st.error("Failed to analyze text. Please try again.")
#                 return None
                
#             progress_bar.progress(70)
#             status_text.text('Document content generated successfully!')
            
#             # Upload and log
#             status_text.text('Uploading to HackMD...')
#             hackmd_link = self.upload_and_log(analyzed_text, document_type)
            
#             progress_bar.progress(100)
#             status_text.text('Process completed successfully!')
            
#             return analyzed_text
            
#         except Exception as e:
#             logger.error(f"Error generating assignment document: {str(e)}")
#             st.error(f"Error generating document: {str(e)}")
#             return None


# class ConceptsDocumentGenerator(DocumentGenerator):
#     """Generator for Concepts Reference Documents with rate limit handling"""
    
#     def generate(self, uploaded_files: List) -> Optional[str]:
#         """Generate Concepts Reference Document with rate limit handling"""
#         if not self._validate_parameters():
#             return None
            
#         document_type = "Concepts Reference Document"
        
#         # Progress bar
#         progress_bar = st.progress(0)
#         status_text = st.empty()
        
#         try:
#             # Extract text from all files
#             status_text.text('Reading and analyzing files...')
#             extracted_text = ""
            
#             # Process files in parallel
#             with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
#                 futures = []
#                 for uploaded_file in uploaded_files:
#                     file_type = uploaded_file.name.split(".")[-1].lower()
#                     futures.append(
#                         executor.submit(TextExtractor.extract, uploaded_file, file_type)
#                     )
                
#                 # Collect results as they complete
#                 total_files = len(futures)
#                 for i, future in enumerate(concurrent.futures.as_completed(futures)):
#                     try:
#                         file_text = future.result()
#                         if file_text and file_text.strip():
#                             extracted_text += file_text + "\n"
#                         progress_value = min(30, int(30 * (i + 1) / total_files))
#                         progress_bar.progress(progress_value)
#                     except Exception as e:
#                         logger.error(f"Error processing file {i}: {str(e)}")
#                         st.warning(f"Error processing one of the files: {str(e)}")
            
#             if not extracted_text.strip():
#                 st.error("Failed to extract any text. Please check the uploaded files.")
#                 return None
                
#             status_text.text('Files read successfully!')
            
#             # Create rate limit handler
#             rate_limiter = RateLimitHandler(
#                 requests_per_minute=15,  # Gemini's limit
#                 retry_delay=65,          # Wait a bit over a minute to be safe
#                 max_retries=10           # Allow up to 10 retries
#             )
            
#             # Analyze concepts with rate limiting
#             status_text.text('Analyzing concepts...')
#             prompt = f"""
#             Please follow the steps below for analyzing the extracted content:

#             1. **Analyze Questions from Images**:
#             - First, focus on identifying and extracting any questions from the images. A question is defined as a sentence or phrase that is asking for specific information, typically ending with a question mark.
#             - If the extracted text includes any question-like phrases from images, please isolate them and remove any surrounding context.
#             - Remove unnecessary prefixes like "Asked about," "Can you explain," "What is," or similar phrases. 
#             - **If there is any question or concept related to 'Introduce Yourself,' remove it entirely from the analysis.**
#             - Ignore any broad or generic questions such as:
#                 - Questions about educational background (e.g., "What was your major technology focus during BE?", "Technologies Used In Current Projects")
#                 - Questions about personal details or self-introduction (e.g., "Can you introduce yourself?")
#                 - Resume-related questions (e.g., "Rate your skills as mentioned on your resume")
#                 - Generic or vague questions that do not focus on specific

#             2. **Extract Key Concepts**:
#             - Once the questions have been analyzed and extracted, identify the key concepts or topics they are related to. Each concept should be a descriptive phrase that directly reflects the essence of the question. Avoid using overly generic terms.
#             - Use longer, descriptive names for concepts to ensure clarity and context. For example:
#                 - Instead of "Polymorphism," use "Polymorphism in Object-Oriented Programming."
#                 - Instead of "HTTP," use "HTTP Protocol and Its Methods."

#             3. **Filter Unique Concepts**:
#             - Remove duplicate concept names. Display only unique, descriptive concepts, ensuring there are no repeated entries in the list.

#             4. *Add Similar Concepts*:
#             - Based on the extracted unique concepts, generate 10 additional descriptive and relevant concepts aligned with the context.

#             5. **Provide Preview**:
#             - Display the list of unique extracted concepts followed by the list of similar concepts. These should be displayed as a simple list with each concept on a new line and each concept name starting with a capital letter.
#             - **Do not include any extra context, such as headers like "Key Concepts" or "Extracted Questions." Just show the concept names cleanly.**

#             Here is the content to analyze:

#             {extracted_text}
#             """
            
#             # Use rate limiter for API call
#             analyzed_concepts = rate_limiter.execute_with_rate_limit(analyze_text_with_gemini, prompt)
            
#             if not analyzed_concepts or analyzed_concepts.strip() == "":
#                 st.error("Failed to analyze concepts. Please verify the content.")
#                 return None
                
#             progress_bar.progress(50)
#             status_text.text('Concepts analyzed successfully!')
            
#             # Display concepts preview
#             concepts_list = analyzed_concepts.split("\n")
#             filtered_concepts = [
#                 concept.strip() 
#                 for concept in concepts_list 
#                 if concept.strip() and "?" not in concept
#             ]
            
#             # Show the number of concepts found
#             num_concepts = len(filtered_concepts)
#             if num_concepts > 15:
#                 st.warning(f"Found {num_concepts} concepts. This may exceed Gemini's rate limit of 15 requests per minute. The process will automatically pause when needed and resume after the rate limit resets.")
            
#             st.markdown("### Preview of Extracted Concepts:")
#             preview_container = st.container()
#             with preview_container:
#                 for concept in filtered_concepts:
#                     st.write(f"- {concept}")
            
#             # Option to select only specific concepts for processing
#             if len(filtered_concepts) > 5:
#                 st.info("If you'd like to process only specific concepts, you can select them below. Otherwise, all concepts will be processed.")
#                 selected_concepts = st.multiselect(
#                     "Select specific concepts to process (optional)",
#                     filtered_concepts,
#                     default=None
#                 )
                
#                 if selected_concepts:
#                     filtered_concepts = selected_concepts
#                     st.success(f"Processing {len(filtered_concepts)} selected concepts.")
            
#             # Generate detailed content for each concept
#             status_text.text('Generating detailed explanations for concepts...')
#             full_document_content = ""
            
#             # Store already processed concepts in case of interruption
#             processed_concepts = {}
#             concept_status = st.empty()
            
#             # Create placeholder for rate limit info
#             rate_limit_info = st.empty()
            
#             # Process concepts sequentially with rate limiting
#             for i, concept in enumerate(filtered_concepts):
#                 concept_number = i + 1
#                 concept_status.info(f"Processing concept {concept_number}/{len(filtered_concepts)}: {concept}")
                
#                 solution_prompt = f"""
#                 Provide a theoretical answer for the concept: {concept}, ensuring key details are covered without including a summary or conclusion.

#                 **For Coding-Related Questions**:
#                 - Do NOT provide complete code.
#                 - Instead, describe the approach in a structured manner:
#                 - Explain the problem briefly.
#                 - Outline the key steps needed to solve it.
#                 - Mention important considerations like time complexity, edge cases, and best practices.

#                 **For Non-Coding Theoretical Concepts**:
#                 - Provide an in-depth explanation covering relevant details.
#                 - Use bullet points or structured formatting where appropriate.

#                 Concept: {concept}
#                 """
                
#                 try:
#                     # Use rate limiter for API call
#                     solution = rate_limiter.execute_with_rate_limit(analyze_text_with_gemini, solution_prompt)
                    
#                     if solution:
#                         processed_concepts[concept] = solution
#                         full_document_content += f"### {concept}\n\n{solution}\n\n"
                        
#                     # Update progress
#                     progress_value = min(80, 50 + int(30 * (i + 1) / len(filtered_concepts)))
#                     progress_bar.progress(progress_value)
#                     status_text.text(f'Generated {i+1}/{len(filtered_concepts)} concepts...')
                    
#                     # Show remaining rate limit info
#                     remaining_requests = 15 - len(rate_limiter.request_timestamps)
#                     rate_limit_info.info(f"API calls remaining before rate limit: {max(0, remaining_requests)}/15")
                    
#                 except Exception as e:
#                     logger.error(f"Error generating solution for concept '{concept}': {str(e)}")
#                     st.warning(f"Error generating solution for concept '{concept}': {str(e)}")
#                     # Continue with next concept rather than stopping completely
#                     continue
            
#             # Upload and log
#             if full_document_content:
#                 status_text.text('Uploading to HackMD...')
#                 hackmd_link = self.upload_and_log(full_document_content, document_type)
                
#                 progress_bar.progress(100)
#                 status_text.text('Process completed successfully!')
                
#                 # Clear temporary status indicators
#                 concept_status.empty()
#                 rate_limit_info.empty()
                
#                 return full_document_content
#             else:
#                 st.error("Failed to generate any content for concepts. Please try again.")
#                 return None
                
#         except Exception as e:
#             logger.error(f"Error generating concepts document: {str(e)}")
#             st.error(f"Error generating document: {str(e)}")
#             return None


# class BatchConceptsDocumentGenerator(DocumentGenerator):
#     """Generator for Concepts Reference Documents using batch processing"""
    
#     def generate(self, uploaded_files: List) -> Optional[str]:
#         """Generate Concepts Reference Document with batch processing"""
#         if not self._validate_parameters():
#             return None
            
#         document_type = "Concepts Reference Document"
        
#         # Progress bar
#         progress_bar = st.progress(0)
#         status_text = st.empty()
        
#         try:
#             # Extract text from all files - this part remains unchanged
#             status_text.text('Reading and analyzing files...')
#             extracted_text = ""
            
#             with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
#                 futures = []
#                 for uploaded_file in uploaded_files:
#                     file_type = uploaded_file.name.split(".")[-1].lower()
#                     futures.append(
#                         executor.submit(TextExtractor.extract, uploaded_file, file_type)
#                     )
                
#                 total_files = len(futures)
#                 for i, future in enumerate(concurrent.futures.as_completed(futures)):
#                     try:
#                         file_text = future.result()
#                         if file_text and file_text.strip():
#                             extracted_text += file_text + "\n"
#                         progress_value = min(30, int(30 * (i + 1) / total_files))
#                         progress_bar.progress(progress_value)
#                     except Exception as e:
#                         logger.error(f"Error processing file {i}: {str(e)}")
#                         st.warning(f"Error processing one of the files: {str(e)}")
            
#             if not extracted_text.strip():
#                 st.error("Failed to extract any text. Please check the uploaded files.")
#                 return None
                
#             status_text.text('Files read successfully!')
            
#             # Analyze concepts
#             status_text.text('Analyzing concepts...')
#             prompt = f"""
#             Please follow the steps below for analyzing the extracted content:

#             1. **Analyze Questions from Images**:
#             - First, focus on identifying and extracting any questions from the images. A question is defined as a sentence or phrase that is asking for specific information, typically ending with a question mark.
#             - If the extracted text includes any question-like phrases from images, please isolate them and remove any surrounding context.
#             - Remove unnecessary prefixes like "Asked about," "Can you explain," "What is," or similar phrases. 
#             - **If there is any question or concept related to 'Introduce Yourself,' remove it entirely from the analysis.**
#             - Ignore any broad or generic questions such as:
#                 - Questions about educational background (e.g., "What was your major technology focus during BE?", "Technologies Used In Current Projects")
#                 - Questions about personal details or self-introduction (e.g., "Can you introduce yourself?")
#                 - Resume-related questions (e.g., "Rate your skills as mentioned on your resume")
#                 - Generic or vague questions that do not focus on specific

#             2. **Extract Key Concepts**:
#             - Once the questions have been analyzed and extracted, identify the key concepts or topics they are related to. Each concept should be a descriptive phrase that directly reflects the essence of the question. Avoid using overly generic terms.
#             - Use longer, descriptive names for concepts to ensure clarity and context. For example:
#                 - Instead of "Polymorphism," use "Polymorphism in Object-Oriented Programming."
#                 - Instead of "HTTP," use "HTTP Protocol and Its Methods."

#             3. **Filter Unique Concepts**:
#             - Remove duplicate concept names. Display only unique, descriptive concepts, ensuring there are no repeated entries in the list.

#             4. *Add Similar Concepts*:
#             - Based on the extracted unique concepts, generate 10 additional descriptive and relevant concepts aligned with the context.

#             5. **Provide Preview**:
#             - Display the list of unique extracted concepts followed by the list of similar concepts. These should be displayed as a simple list with each concept on a new line and each concept name starting with a capital letter.
#             - **Do not include any extra context, such as headers like "Key Concepts" or "Extracted Questions." Just show the concept names cleanly.**

#             Here is the content to analyze:

#             {extracted_text}
#             """
            
#             # We'll still make a single call for the concept extraction
#             analyzed_concepts = analyze_text_with_gemini(prompt)
            
#             if not analyzed_concepts or analyzed_concepts.strip() == "":
#                 st.error("Failed to analyze concepts. Please verify the content.")
#                 return None
                
#             progress_bar.progress(40)
#             status_text.text('Concepts analyzed successfully!')
            
#             # Display concepts preview
#             concepts_list = analyzed_concepts.split("\n")
#             filtered_concepts = [
#                 concept.strip() 
#                 for concept in concepts_list 
#                 if concept.strip() and "?" not in concept
#             ]
            
#             # Show the number of concepts found
#             num_concepts = len(filtered_concepts)
#             if num_concepts > 15:
#                 st.warning(f"Found {num_concepts} concepts. Using batch processing to handle them efficiently.")
            
#             st.markdown("### Preview of Extracted Concepts:")
#             for concept in filtered_concepts:
#                 st.write(f"- {concept}")
                
#             # Option to select only specific concepts for processing
#             if len(filtered_concepts) > 5:
#                 st.info("You can select specific concepts to process or process all of them.")
                
#                 # Option to process all concepts or select specific ones
#                 process_mode = st.radio(
#                     "Select processing mode:",
#                     ["Process all concepts", "Select specific concepts"]
#                 )
                
#                 if process_mode == "Select specific concepts":
#                     selected_concepts = st.multiselect(
#                         "Select concepts to process",
#                         filtered_concepts,
#                         default=filtered_concepts[:5]  # Default to first 5
#                     )
                    
#                     if selected_concepts:
#                         filtered_concepts = selected_concepts
#                         st.success(f"Processing {len(filtered_concepts)} selected concepts.")
            
#             # Ask for batch size
#             col1, col2 = st.columns(2)
#             with col1:
#                 batch_size = st.number_input(
#                     "Concepts per batch (recommended: 10-15 for Gemini free tier)",
#                     min_value=1,
#                     max_value=50,
#                     value=10
#                 )
            
#             with col2:
#                 batch_delay = st.number_input(
#                     "Seconds to wait between batches",
#                     min_value=30,
#                     max_value=120,
#                     value=65
#                 )
            
#             # Initialize batch processor
#             def update_progress(current, total):
#                 progress_value = min(80, 40 + int(40 * current / total))
#                 progress_bar.progress(progress_value)
#                 status_text.text(f'Generated {current}/{total} concepts...')
            
#             batch_processor = BatchConceptsProcessor(
#                 batch_size=batch_size,
#                 batch_delay=batch_delay,
#                 api_function=analyze_text_with_gemini,
#                 progress_callback=update_progress
#             )
            
#             if st.button("Start Processing Concepts"):
#                 # Generate detailed content for each concept using batch processing
#                 status_text.text('Generating detailed explanations for concepts...')
                
#                 # Process concepts in batches
#                 concept_results = batch_processor.process_concepts(
#                     concepts=filtered_concepts,
#                     prompt_template=solution_prompt_for_concept
#                 )
                
#                 # Generate the full document content
#                 full_document_content = ""
                
#                 for concept in filtered_concepts:
#                     if concept in concept_results:
#                         solution = concept_results[concept]
#                         full_document_content += f"### {concept}\n\n{solution}\n\n"
                
#                 # Upload and log
#                 if full_document_content:
#                     status_text.text('Uploading to HackMD...')
#                     hackmd_link = self.upload_and_log(full_document_content, document_type)
                    
#                     progress_bar.progress(100)
#                     status_text.text('Process completed successfully!')
                    
#                     return full_document_content
#                 else:
#                     st.error("Failed to generate any content for concepts. Please try again.")
#                     return None
            
#             return None
                
#         except Exception as e:
#             logger.error(f"Error generating concepts document: {str(e)}")
#             st.error(f"Error generating document: {str(e)}")
#             return None


# def generate_reference_document_to_docx(analyzed_text, output_filename_docx, output_stream):
#     """Generate a formatted DOCX document from markdown text"""
#     doc = Document()

#     paragraphs = analyzed_text.split("\n\n")  

#     for paragraph in paragraphs:
#         lines = paragraph.splitlines()
#         for line in lines:
#             if line.startswith('#### '):
#                 p = doc.add_paragraph(line[5:])
#                 p.style = 'Heading 4'
#             elif line.startswith('### '):
#                 p = doc.add_paragraph(line[4:])
#                 p.style = 'Heading 3'
#             elif line.startswith('## '):
#                 p = doc.add_paragraph(line[3:])
#                 p.style = 'Heading 2'
#             elif line.startswith('# '):
#                 p = doc.add_paragraph(line[2:])
#                 p.style = 'Heading 1'
#             else:
#                 p = doc.add_paragraph()
#                 parts = re.split(r'(\*\*.*?\*\*)', line)
#                 for part in parts:
#                     if part.startswith('**') and part.endswith('**'):
#                         run = p.add_run(part[2:-2])  
#                         run.bold = True
#                     else:
#                         p.add_run(part)

#     doc.save(output_stream)


# def show_usage_instructions():
#     """Show instructions for using the app"""
#     with st.expander("How to use this app"):
#         st.markdown("""
#         ### Assignment Reference Document Generator
        
#         This tool helps you generate reference documents for assignments:
        
#         1. **Select document type**: Choose "Assignment Reference Document"
#         2. **Fill in the details**: Title, Job ID, your name, and job type
#         3. **Upload file**: Upload a PDF, DOCX, or text file containing the assignment
#         4. **Generate**: Click the "Generate Document" button
#         5. **Get link**: The tool will generate a document and provide a HackMD link
        
#         ### Concepts Reference Document Generator
        
#         This tool analyzes interview questions and generates explanations:
        
#         1. **Select document type**: Choose "Concepts Reference Document"
#         2. **Choose processing mode**: Standard or Batch processing
#         3. **Fill in the details**: Title, Job ID, your name, job type, and interview round
#         4. **Upload files**: Upload files containing interview questions/concepts
#         5. **Generate**: Click the "Generate Document" button
#         6. **Review concepts**: The tool will extract concepts for your review
#         7. **Get link**: A HackMD link with all explanations will be provided
#         """)


# def reference_document_generator():
#     """Main Streamlit UI function for the Reference Document Generator"""
#     st.header("Reference Document Generator")
    
#     # Add styling
#     st.markdown("""
#     <style>
#     .stButton button {
#         background-color: #4CAF50;
#         color: white;
#         padding: 10px 20px;
#         border: none;
#         border-radius: 4px;
#         cursor: pointer;
#     }
#     .stButton button:hover {
#         background-color: #45a049;
#     }
#     .warning {
#         color: red;
#         font-weight: bold;
#     }
#     </style>
#     """, unsafe_allow_html=True)

#     # App tabs
#     tab1, tab2 = st.tabs(["Generate Document", "Download Previous Document"])
    
#     with tab1:
#         # Document type selection
#         document_type = st.selectbox("Select Document Type", ["Assignment Reference Document", "Concepts Reference Document"])

#         # For Concepts Reference Document, add processing mode selection
#         processing_mode = None
#         if document_type == "Concepts Reference Document":
#             processing_mode = st.radio(
#                 "Select processing mode:",
#                 ["Standard (Automatic pause and retry)", "Batch processing (Manual control)"]
#             )

#         # Input fields with validation
#         with st.form(key="document_form"):
#             title_name = st.text_input("Enter Title Name for the Document", 
#                                     placeholder="Enter the title to be displayed in HackMD",
#                                     help="This will be displayed as the title in HackMD")
            
#             job_id = st.text_input("Enter Job ID", 
#                                 placeholder="Enter the Job ID associated with this document",
#                                 help="This will be used for tracking in the Google Sheet")
            
#             prepared_by = st.text_input("Prepared By", 
#                                     placeholder="Enter your name",
#                                     help="Your name will be recorded in the logs")
            
#             job_type = st.selectbox("Select Job Type", 
#                                 ["Full Time", "Internships", "Intern+FullTime"],
#                                 help="Type of job this document is for")

#             # Show Interview Round dropdown for "Concepts Reference Document"
#             interview_round = None
#             if document_type == "Concepts Reference Document":
#                 interview_round = st.selectbox(
#                     "Select Interview Round",
#                     [
#                         "Assessment", "TR1", "TR2", "MR", "HR", 
#                         "TR1 + TR2 + HR", "Assessment + TR1 + TR2 + HR", 
#                         "Offline drive", "CEO Round", "Culture Fit Round"
#                     ],
#                     help="The interview round this document is for"
#                 )

#             # File Uploader with clear instructions
#             st.write("### Upload Files")
            
#             if document_type == "Assignment Reference Document":
#                 st.write("Please upload the assignment document (PDF, DOCX, or text file)")
#                 uploaded_files = st.file_uploader(
#                     "Upload assignment document", 
#                     type=["pdf", "docx", "txt"], 
#                     accept_multiple_files=False,
#                     help="Upload a single file containing the assignment description"
#                 )
#             else:  # Concepts Reference Document
#                 st.write("Please upload interview questions or concepts (PDF, DOCX, ZIP, Image files)")
#                 st.write("For best results, ensure images are clear and readable")
#                 uploaded_files = st.file_uploader(
#                     "Upload files", 
#                     type=SUPPORTED_FILE_TYPES, 
#                     accept_multiple_files=True,
#                     help="Upload one or more files containing interview questions or concepts"
#                 )
            
#             submit_button = st.form_submit_button("Generate Document")
        
#         # Process form submission
#         if submit_button:
#             # Basic validation
#             if not uploaded_files:
#                 st.error("Please upload at least one file before proceeding.")
#                 st.stop()
                
#             # For Assignment Reference Document
#             if document_type == "Assignment Reference Document":
#                 generator = AssignmentDocumentGenerator(
#                     title_name, 
#                     job_id,
#                     prepared_by,
#                     job_type
#                 )
#                 generator.generate(uploaded_files)
                
#             # For Concepts Reference Document
#             elif document_type == "Concepts Reference Document":
#                 if not isinstance(uploaded_files, list):
#                     uploaded_files = [uploaded_files]
                
#                 # Choose generator based on processing mode
#                 if processing_mode == "Standard (Automatic pause and retry)":
#                     generator = ConceptsDocumentGenerator(
#                         title_name,
#                         job_id,
#                         prepared_by,
#                         job_type,
#                         interview_round
#                     )
#                 else:  # Batch processing
#                     generator = BatchConceptsDocumentGenerator(
#                         title_name,
#                         job_id,
#                         prepared_by,
#                         job_type,
#                         interview_round
#                     )
                
#                 generator.generate(uploaded_files)
    
#     with tab2:
#         st.write("### Download Previously Generated Document")
#         st.write("Enter the HackMD link of the document you want to download:")
        
#         hackmd_url = st.text_input("HackMD URL", placeholder="https://hackmd.io/...")
        
#         if st.button("Download as DOCX") and hackmd_url:
#             with st.spinner("Fetching document and converting to DOCX..."):
#                 try:
#                     # Extract document ID from URL
#                     doc_id = hackmd_url.split("/")[-1]
                    
#                     # Fetch content from HackMD API
#                     api_url = f"https://api.hackmd.io/v1/notes/{doc_id}"
#                     headers = {"Authorization": f"Bearer {HACKMD_API_TOKEN}"}
                    
#                     response = requests.get(api_url, headers=headers)
#                     if response.status_code == 200:
#                         content = response.json().get("content", "")
#                         title = response.json().get("title", "Document")
                        
#                         # Generate DOCX
#                         output = BytesIO()
#                         generate_reference_document_to_docx(content, f"{title}.docx", output)
#                         output.seek(0)
                        
#                         # Offer download
#                         st.download_button(
#                             label="Download DOCX",
#                             data=output,
#                             file_name=f"{title}.docx",
#                             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#                         )
#                     else:
#                         st.error(f"Failed to fetch document. Status code: {response.status_code}")
#                 except Exception as e:
#                     logger.error(f"Error downloading document: {str(e)}")
#                     st.error(f"Error downloading document: {str(e)}")


# def main():
#     """Main function to run the Streamlit app"""
#     # Configure page settings
#     st.set_page_config(
#         page_title="Reference Document Generator",
#         page_icon="📄",
#         layout="wide",
#         initial_sidebar_state="expanded"
#     )
    
#     # App header with logo and title
#     col1, col2 = st.columns([1, 5])
#     with col1:
#         st.image("https://res.cloudinary.com/dg8n2jeur/image/upload/v1741718114/cws2futayeddahz3sz8z.webp", width=80)
#     with col2:
#         st.title("Reference Document Generator")
#         st.markdown("Create professional reference documents for assignments and interview concepts")
    
#     # Sidebar information
#     with st.sidebar:
#         st.header("About")
#         st.info(
#             "This application helps create structured reference documents "
#             "for assignments and technical interview concepts. Documents are "
#             "automatically uploaded to HackMD and logged in Google Sheets."
#         )
        
#         st.header("Options")
#         # Add theme selector or other preferences here
#         theme = st.selectbox("Theme", ["Light", "Dark"])
#         if theme == "Dark":
#             # Apply dark theme CSS
#             st.markdown("""
#             <style>
#             body {
#                 background-color: #121212;
#                 color: #ffffff;
#             }
#             </style>
#             """, unsafe_allow_html=True)
        
#         # Show rate limit information
#         st.header("API Rate Limit Info")
#         st.warning(
#             "Gemini API has a rate limit of 15 requests per minute. "
#             "When generating solutions for many concepts, the application "
#             "will automatically pause when needed and resume after the rate limit resets."
#         )
        
#         st.info(
#             "For large numbers of concepts, consider using the 'Batch processing' mode "
#             "which allows you to configure batch sizes and processing intervals."
#         )
        

#     # Check dependencies
#     dependency_error = False
#     try:
#         import gspread
#         import requests
#         import pytesseract
#         import google.generativeai
#     except ImportError as e:
#         st.error(f"Missing dependency: {str(e)}")
#         dependency_error = True
    
#     if dependency_error:
#         st.warning("Some features may not work properly due to missing dependencies.")
    
#     # Show instruction expander
#     show_usage_instructions()
    
#     # Main app
#     reference_document_generator()
    
#     # Footer
#     st.markdown("---")
#     st.markdown(
#         "© 2025 Reference Document Generator | "
#         "For support, contact tirupathirao.kella@nxtwave.co.in"
#     )


# if __name__ == "__main__":
#     # Set up exception handling
#     try:
#         main()
#     except Exception as e:
#         st.error(f"An unexpected error occurred: {str(e)}")
#         logger.critical(f"Critical application error: {str(e)}", exc_info=True)



def analyze_text_with_gemini(prompt):
    try:
        model = genai.GenerativeModel("gemini-2.0-flash")  # Use the appropriate Gemini model
        response = model.generate_content(prompt)
        return response.text.strip() if response.text else "No response generated."
    except Exception as e:
        return f"Error analyzing text: {str(e)}"

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Import utilities with error handling
try:
    from utils.pdf_processing import extract_text_from_pdf
except ImportError as e:
    logger.error(f"Failed to import required modules: {str(e)}")
    st.error(f"Failed to import required modules: {str(e)}")
    
# Constants
SPREADSHEET_ID = "1EQiw3jPnTamigx9ZHqOKr6zLl9FlQmko7IQnjlfRtqo"
CREDENTIALS_PATH = "credentials.json"
HACKMD_API_TOKEN = "1PQ8Z9IMDPD579ETDEMLJYG8W3V2CY6IPLKYVD7Z8MZQW9WPX9"  # Consider moving to secrets
MAX_WORKERS = 4  # For parallel processing
MAX_RETRIES = 3  # For API calls
SUPPORTED_FILE_TYPES = ["pdf", "docx", "zip", "jpg", "jpeg", "png", "txt"]
CACHE_TTL = 3600  # Cache time-to-live in seconds
HACKMD_MAX_CONTENT_LENGTH = 90000  # Reduced size for safer margin (80KB)


class RateLimitHandler:
    """Handles rate limiting for API calls with automatic backoff and retry"""
    
    def __init__(self, 
                requests_per_minute: int = 15,
                retry_delay: int = 65,  # 65 seconds to ensure full minute has passed
                max_retries: int = 5):
        self.requests_per_minute = requests_per_minute
        self.retry_delay = retry_delay
        self.max_retries = max_retries
        self.request_timestamps = []
        
    def _clean_old_timestamps(self):
        """Remove timestamps older than 1 minute"""
        current_time = time.time()
        self.request_timestamps = [ts for ts in self.request_timestamps 
                               if current_time - ts < 60]
        
    def _can_make_request(self) -> bool:
        """Check if we can make a request based on the rate limit"""
        self._clean_old_timestamps()
        return len(self.request_timestamps) < self.requests_per_minute
        
    def _time_until_next_available_slot(self) -> float:
        """Calculate time until we can make another request"""
        if self._can_make_request():
            return 0
            
        self._clean_old_timestamps()
        oldest_timestamp = min(self.request_timestamps)
        return max(0, 60 - (time.time() - oldest_timestamp))
        
    def execute_with_rate_limit(self, 
                              func: Callable, 
                              *args, 
                              **kwargs) -> Any:
        """
        Execute a function with rate limiting
        Returns the function result or raises the last exception after max retries
        """
        retries = 0
        last_exception = None
        
        while retries <= self.max_retries:
            try:
                # Check if we can make a request
                wait_time = self._time_until_next_available_slot()
                
                if wait_time > 0:
                    # Need to wait before making another request
                    retries += 1
                    if retries > self.max_retries:
                        break
                        
                    st.warning(f"Rate limit reached. Waiting {int(wait_time)} seconds before continuing...")
                    time.sleep(wait_time)
                    continue
                
                # Record this request timestamp
                self.request_timestamps.append(time.time())
                
                # Execute the function
                return func(*args, **kwargs)
                
            except Exception as e:
                last_exception = e
                error_message = str(e).lower()
                
                # Check if it's a rate limit error
                if "quota" in error_message or "rate" in error_message or "limit" in error_message:
                    retries += 1
                    if retries > self.max_retries:
                        break
                        
                    st.warning(f"API rate limit exceeded. Waiting {self.retry_delay} seconds before retry {retries}/{self.max_retries}...")
                    time.sleep(self.retry_delay)
                else:
                    # If it's not a rate limit error, re-raise immediately
                    raise
        
        # If we get here, we've exceeded max retries
        error_msg = f"Failed after {self.max_retries} retries: {str(last_exception)}"
        st.error(error_msg)
        raise Exception(error_msg)


class BatchConceptsProcessor:
    """Processes concepts in batches to manage API rate limits"""
    
    def __init__(self, 
                batch_size: int = 10, 
                batch_delay: int = 65,
                api_function: Callable = None,
                progress_callback: Callable = None):
        self.batch_size = batch_size
        self.batch_delay = batch_delay
        self.api_function = api_function
        self.progress_callback = progress_callback
        self.results = {}
        self.errors = {}
        
    def process_concepts(self, 
                       concepts: List[str], 
                       prompt_template: str) -> Dict[str, str]:
        """Process concepts in batches with rate limit handling"""
        
        # Create batches of concepts
        batches = []
        current_batch = []
        
        for concept in concepts:
            current_batch.append(concept)
            if len(current_batch) >= self.batch_size:
                batches.append(current_batch)
                current_batch = []
                
        # Add the last batch if it's not empty
        if current_batch:
            batches.append(current_batch)
            
        # Process each batch
        total_processed = 0
        
        # Create a status container
        status_container = st.container()
        
        for batch_idx, batch in enumerate(batches):
            with status_container:
                st.info(f"Processing batch {batch_idx+1}/{len(batches)} ({len(batch)} concepts)")
                
                batch_progress = st.progress(0)
                
                # Process each concept in the batch
                for i, concept in enumerate(batch):
                    # Format the prompt for this concept
                    prompt = prompt_template.format(concept=concept)
                    
                    try:
                        # Call the API function
                        result = self.api_function(prompt)
                        self.results[concept] = result
                    except Exception as e:
                        error_msg = str(e)
                        logger.error(f"Error processing concept '{concept}': {error_msg}")
                        self.errors[concept] = error_msg
                        st.warning(f"Error processing concept: {concept}")
                        
                    # Update progress
                    total_processed += 1
                    batch_progress.progress((i + 1) / len(batch))
                    
                    if self.progress_callback:
                        self.progress_callback(total_processed, len(concepts))
                
                # Reset the batch progress
                batch_progress.empty()
                
            # Wait between batches to avoid rate limits
            if batch_idx < len(batches) - 1:  # Don't wait after the last batch
                with status_container:
                    wait_text = st.text(f"Waiting {self.batch_delay} seconds for rate limit to reset...")
                    wait_progress = st.progress(0)
                    
                    # Show a countdown timer
                    for second in range(self.batch_delay):
                        wait_progress.progress((second + 1) / self.batch_delay)
                        wait_text.text(f"Waiting for rate limit to reset... {self.batch_delay - second} seconds remaining")
                        time.sleep(1)
                    
                    wait_text.empty()
                    wait_progress.empty()
        
        # Return all processed results
        return self.results


def solution_prompt_for_concept(concept: str) -> str:
    """Generate a solution prompt for a concept"""
    return f"""
    Provide a theoretical answer for the concept: {concept}, ensuring key details are covered without including a summary or conclusion.

    **For Coding-Related Questions**:
    - Do NOT provide complete code.
    - Instead, describe the approach in a structured manner:
    - Explain the problem briefly.
    - Outline the key steps needed to solve it.
    - Mention important considerations like time complexity, edge cases, and best practices.

    **For Non-Coding Theoretical Concepts**:
    - Provide an in-depth explanation covering relevant details.
    - Use bullet points or structured formatting where appropriate.

    Concept: {concept}
    """


class GoogleSheetLogger:
    """Class to handle Google Sheets logging operations"""
    
    def __init__(self, spreadsheet_id: str):
        self.spreadsheet_id = spreadsheet_id
        self._gc = None
        
    @property
    def gc(self):
        """Lazy loading of Google Sheets client"""
        if self._gc is None:
            try:
                # Load credentials from Streamlit secrets
                credentials = st.secrets["google_service_account"]
                self._gc = gspread.service_account_from_dict(credentials)
            except Exception as e:
                logger.error(f"Failed to initialize Google Sheets client: {str(e)}")
                st.error(f"Failed to initialize Google Sheets client: {str(e)}")
                raise
        return self._gc
    
    def log_document(self, 
                    job_type: str, 
                    created_at: str, 
                    prepared_by: str, 
                    job_id: str, 
                    document_type: str, 
                    title_name: str, 
                    hackmd_link: str, 
                    interview_round: Optional[str] = None) -> bool:
        """Log document information to Google Sheets"""
        try:
            # Open the spreadsheet and select the "Reference Docs" sheet
            sheet = self.gc.open_by_key(self.spreadsheet_id).worksheet("Reference Docs")

            # Prepare the row to include all fields, adding interview_round if present
            row_data = [job_type, created_at, prepared_by, job_id, document_type, title_name, hackmd_link]
            if interview_round:
                row_data.append(interview_round)

            # Append the row to the specific sheet
            sheet.append_row(row_data)
            logger.info(f"Successfully logged document: {title_name}")
            return True
        except Exception as e:
            logger.error(f"Failed to log data to Google Sheets: {str(e)}")
            st.error(f"Failed to log data to Google Sheets: {str(e)}")
            return False


class HackMDUploader:
    """Class to handle HackMD upload operations with content splitting"""
    
    def __init__(self, api_token: str):
        self.api_token = api_token
        self.headers = {
            "Authorization": f"Bearer {self.api_token}",
            "Content-Type": "application/json"
        }
        self.max_content_length = HACKMD_MAX_CONTENT_LENGTH
        
    @staticmethod
    def _add_retry_mechanism(func):
        """Decorator to add retry mechanism to API calls"""
        def wrapper(*args, **kwargs):
            for attempt in range(MAX_RETRIES):
                try:
                    return func(*args, **kwargs)
                except requests.RequestException as e:
                    if attempt < MAX_RETRIES - 1:
                        wait_time = 2 ** attempt  # Exponential backoff
                        logger.warning(f"Attempt {attempt+1} failed: {str(e)}. Retrying in {wait_time}s...")
                        time.sleep(wait_time)
                    else:
                        logger.error(f"All {MAX_RETRIES} attempts failed: {str(e)}")
                        raise
        return wrapper
    
    @_add_retry_mechanism
    def _create_single_note(self, content: str, title: str) -> Optional[str]:
        """Create a single HackMD note and return its URL"""
        try:
            create_url = "https://api.hackmd.io/v1/notes"
            
            data = {
                "content": content,
                "title": title,
                "readPermission": "guest",
                "writePermission": "owner",
                "publishType": "freely",
                "commentPermission": "disabled"
            }

            response = requests.post(create_url, json=data, headers=self.headers)

            if response.status_code == 201:
                note_data = response.json()
                publish_link = note_data.get("publishLink", None)
                
                if publish_link:
                    logger.info(f"Successfully uploaded to HackMD: {title}")
                    return publish_link
                else:
                    logger.error("Publish link not found in the response.")
                    return None
            else:
                logger.error(f"Failed to create note. Status Code: {response.status_code}, Response: {response.text}")
                return None
                
        except Exception as e:
            logger.error(f"An error occurred while uploading to HackMD: {str(e)}")
            return None
    
    def _split_content(self, content: str) -> List[str]:
        """Split content into chunks of acceptable size with improved handling for large sections"""
        # Reduce the maximum chunk size to be safer
        safe_max_size = int(self.max_content_length * 0.8)  # 80% of limit as safety margin
        
        chunks = []
        current_chunk = ""
        
        # First try splitting by paragraphs (standard approach)
        paragraphs = content.split("\n\n")
        
        for paragraph in paragraphs:
            # If a single paragraph is too large, we need to split it further
            if len(paragraph) > safe_max_size:
                # If this is a large section, split it by lines
                lines = paragraph.split("\n")
                
                # Reset the current sub-chunk
                sub_chunk = ""
                
                for line in lines:
                    # If adding this line would exceed the limit, start a new chunk
                    if len(sub_chunk) + len(line) + 2 > safe_max_size:
                        if sub_chunk:
                            chunks.append(sub_chunk)
                        sub_chunk = line + "\n"
                    else:
                        sub_chunk += line + "\n"
                
                # Add the last sub-chunk if it's not empty
                if sub_chunk:
                    chunks.append(sub_chunk)
            # Normal case - add paragraph if it fits, otherwise start new chunk
            elif len(current_chunk) + len(paragraph) + 2 > safe_max_size:
                chunks.append(current_chunk)
                current_chunk = paragraph + "\n\n"
            else:
                current_chunk += paragraph + "\n\n"
        
        # Add the last chunk if it's not empty
        if current_chunk:
            chunks.append(current_chunk)
        
        # Final safety check - ensure no chunk exceeds the limit
        final_chunks = []
        for chunk in chunks:
            if len(chunk) > safe_max_size:
                # Split by newlines if still too large
                lines = chunk.split("\n")
                current_subchunk = ""
                
                for line in lines:
                    if len(current_subchunk) + len(line) + 1 > safe_max_size:
                        final_chunks.append(current_subchunk)
                        current_subchunk = line + "\n"
                    else:
                        current_subchunk += line + "\n"
                        
                if current_subchunk:
                    final_chunks.append(current_subchunk)
            else:
                final_chunks.append(chunk)
                
        return final_chunks
    
    @_add_retry_mechanism
    def _get_note_content(self, note_url: str) -> Optional[str]:
        """Get the content of a HackMD note by its URL"""
        try:
            # Extract note ID from URL
            note_id = note_url.split("/")[-1]
            
            # Fetch note content
            api_url = f"https://api.hackmd.io/v1/notes/{note_id}"
            response = requests.get(api_url, headers=self.headers)
            
            if response.status_code == 200:
                return response.json().get("content")
            else:
                logger.error(f"Failed to get note content. Status Code: {response.status_code}")
                return None
        except Exception as e:
            logger.error(f"Error getting note content: {str(e)}")
            return None
            
    @_add_retry_mechanism
    def _update_note_content(self, note_url: str, new_content: str) -> bool:
        """Update the content of a HackMD note with improved handling for 202 responses"""
        try:
            # Extract note ID from URL
            note_id = note_url.split("/")[-1]
            
            # Update note content
            api_url = f"https://api.hackmd.io/v1/notes/{note_id}"
            data = {"content": new_content}
            
            response = requests.patch(api_url, json=data, headers=self.headers)
            
            if response.status_code == 200:
                logger.info(f"Successfully updated note: {note_url}")
                return True
            elif response.status_code == 202:
                # Status code 202 means "Accepted" - the request is being processed
                logger.info(f"Update request accepted, waiting for processing: {note_url}")
                
                # Wait for the update to complete
                time.sleep(3)  # Wait 3 seconds before considering it successful
                return True
            else:
                logger.error(f"Failed to update note. Status Code: {response.status_code}")
                return False
        except Exception as e:
            logger.error(f"Error updating note content: {str(e)}")
            return False
    
    @_add_retry_mechanism
    def upload(self, content: str, title_name: str) -> Optional[str]:
        """Upload content to HackMD, splitting if necessary, and return the published link"""
        try:
            # Check if content is too long
            if len(content) > self.max_content_length:
                # Split content into multiple parts
                chunks = self._split_content(content)
                
                # Show a progress bar for splitting
                with st.spinner(f"Content is large. Splitting into {len(chunks)} parts..."):
                    part_links = []
                    # Create a simpler index with just the document parts
                    index_content = f"# {title_name}\n\n## Document Parts\n\n"
                    
                    # Create each part
                    for i, chunk in enumerate(chunks):
                        part_title = f"{title_name} - Part {i+1}"
                        part_content = f"# {part_title}\n\n{chunk}"
                        
                        # Upload this part
                        part_link = self._create_single_note(part_content, part_title)
                        
                        if part_link:
                            part_links.append((part_title, part_link))
                            index_content += f"- [{part_title}]({part_link})\n"
                        else:
                            # If upload fails, try with a smaller chunk
                            logger.error(f"Failed to upload part {i+1}")
                            st.error(f"Failed to upload part {i+1}")
                            
                            # Try again with a smaller chunk size
                            if len(chunk) > 10000:  # Only try splitting if the chunk is reasonably large
                                st.warning(f"Trying to split part {i+1} into smaller segments...")
                                
                                # Split in half
                                half_size = len(chunk) // 2
                                first_half = chunk[:half_size]
                                second_half = chunk[half_size:]
                                
                                # Try uploading first half
                                first_half_title = f"{title_name} - Part {i+1}a"
                                first_half_content = f"# {first_half_title}\n\n{first_half}"
                                first_half_link = self._create_single_note(first_half_content, first_half_title)
                                
                                if first_half_link:
                                    part_links.append((first_half_title, first_half_link))
                                    index_content += f"- [{first_half_title}]({first_half_link})\n"
                                
                                # Try uploading second half
                                second_half_title = f"{title_name} - Part {i+1}b"
                                second_half_content = f"# {second_half_title}\n\n{second_half}"
                                second_half_link = self._create_single_note(second_half_content, second_half_title)
                                
                                if second_half_link:
                                    part_links.append((second_half_title, second_half_link))
                                    index_content += f"- [{second_half_title}]({second_half_link})\n"
                    
                    # Create index note with links to all parts - REMOVED the Navigation Links section
                    index_link = self._create_single_note(index_content, title_name)
                    
                    if index_link:
                        st.success(f"Content was split into {len(chunks)} parts due to size limits.")
                        return index_link
                    else:
                        # If index creation fails, return the first part link
                        if part_links:
                            st.warning("Created multiple parts but failed to create index. Using first part as main link.")
                            return part_links[0][1]
                        else:
                            return None
            else:
                # Content is short enough for a single note
                content_with_title = f"# {title_name}\n\n{content}"
                return self._create_single_note(content_with_title, title_name)
                
        except Exception as e:
            logger.error(f"An error occurred while splitting and uploading to HackMD: {str(e)}")
            st.error(f"An error occurred while uploading to HackMD: {str(e)}")
            return None


class TextExtractor:
    """Class to handle text extraction from various file types and direct text input"""
    
    @staticmethod
    def from_direct_input(text: str) -> str:
        """Process directly inputted text"""
        return text
            
    @staticmethod
    def _ensure_tesseract():
        """Check if Tesseract is installed and configured"""
        try:
            pytesseract.get_tesseract_version()
        except Exception as e:
            logger.error(f"Tesseract not properly configured: {str(e)}")
            st.error("Tesseract OCR is not properly configured. Images cannot be processed.")
            return False
        return True
            
    @staticmethod
    def from_pdf(file_obj) -> str:
        """Extract text from PDF file"""
        try:
            return extract_text_from_pdf(file_obj)
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise ValueError(f"Error extracting text from PDF: {str(e)}")
            
    @staticmethod
    def from_docx(file_obj) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_obj)
            return '\n'.join([para.text for para in doc.paragraphs])
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise ValueError(f"Error extracting text from DOCX: {str(e)}")
            
    @staticmethod
    def from_image(file_obj) -> str:
        """Extract text from image using OCR"""
        if not TextExtractor._ensure_tesseract():
            return ""
            
        try:
            img = Image.open(file_obj)
            return pytesseract.image_to_string(img)
        except Exception as e:
            logger.error(f"Error extracting text from image: {str(e)}")
            raise ValueError(f"Error extracting text from image: {str(e)}")
            
    @staticmethod
    def from_zip(file_obj) -> str:
        """Extract text from all supported files in a ZIP archive"""
        if not TextExtractor._ensure_tesseract():
            return ""
            
        text = ""
        try:
            with zipfile.ZipFile(file_obj, "r") as zip_ref:
                for file_name in zip_ref.namelist():
                    ext = file_name.split(".")[-1].lower() if "." in file_name else ""
                    
                    if ext == "txt":
                        with zip_ref.open(file_name) as file:
                            text += file.read().decode('utf-8', errors='replace') + "\n"
                    elif ext in ["jpg", "jpeg", "png"]:
                        with zip_ref.open(file_name) as file:
                            img = Image.open(file)
                            text += pytesseract.image_to_string(img) + "\n"
                    elif ext == "pdf":
                        with zip_ref.open(file_name) as file:
                            # Need to copy to BytesIO since extract_text_from_pdf expects a file-like object
                            pdf_bytes = BytesIO(file.read())
                            text += TextExtractor.from_pdf(pdf_bytes) + "\n"
                    elif ext == "docx":
                        with zip_ref.open(file_name) as file:
                            docx_bytes = BytesIO(file.read())
                            text += TextExtractor.from_docx(docx_bytes) + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting text from ZIP: {str(e)}")
            raise ValueError(f"Error extracting text from ZIP: {str(e)}")
            
    @staticmethod
    def from_txt(file_obj) -> str:
        """Extract text from plain text file"""
        try:
            return file_obj.read().decode('utf-8', errors='replace')
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {str(e)}")
            raise ValueError(f"Error extracting text from TXT: {str(e)}")
            
    @staticmethod
    def extract(file_obj, file_type: str) -> str:
        """Extract text from file based on its type"""
        file_type = file_type.lower()
        
        if file_type not in SUPPORTED_FILE_TYPES and file_type != "direct_input":
            logger.error(f"Unsupported file type: {file_type}")
            raise ValueError(f"Unsupported file type: {file_type}")
            
        extractors = {
            "pdf": TextExtractor.from_pdf,
            "docx": TextExtractor.from_docx,
            "zip": TextExtractor.from_zip,
            "jpg": TextExtractor.from_image,
            "jpeg": TextExtractor.from_image,
            "png": TextExtractor.from_image,
            "txt": TextExtractor.from_txt,
            "direct_input": TextExtractor.from_direct_input
        }
        
        return extractors[file_type](file_obj)


class DocumentGenerator:
    """Base class for document generators"""
    
    def __init__(self, 
                title_name: str, 
                job_id: str, 
                prepared_by: str, 
                job_type: str,
                interview_round: Optional[str] = None):
        self.title_name = title_name
        self.job_id = job_id
        self.prepared_by = prepared_by
        self.job_type = job_type
        self.interview_round = interview_round
        self.gs_logger = GoogleSheetLogger(SPREADSHEET_ID)
        self.hackmd_uploader = HackMDUploader(HACKMD_API_TOKEN)
        
    def _validate_parameters(self) -> bool:
        """Validate input parameters"""
        if not self.title_name.strip():
            st.error("Please enter a valid Title Name before proceeding.")
            return False
            
        if not self.job_id.strip():
            st.error("Please enter a valid Job ID before proceeding.")
            return False
            
        if not self.prepared_by.strip():
            st.error("Please enter the name in 'Prepared By' field before proceeding.")
            return False
            
        if not self.job_type.strip():
            st.error("Please select a valid Job Type before proceeding.")
            return False
            
        return True
        
    def generate(self, *args, **kwargs) -> Optional[str]:
        """Generate document content - to be implemented by subclasses"""
        raise NotImplementedError("Subclasses must implement this method")
        
    def upload_and_log(self, content: str, document_type: str) -> Optional[str]:
        """Upload document to HackMD and log to Google Sheets"""
        if not content:
            logger.error("No content to upload")
            return None
            
        hackmd_link = self.hackmd_uploader.upload(content, self.title_name)
        
        if hackmd_link:
            created_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            self.gs_logger.log_document(
                self.job_type,
                created_at,
                self.prepared_by,
                self.job_id,
                document_type,
                self.title_name,
                hackmd_link,
                self.interview_round
            )
            
            # Display link information in a more user-friendly way
            st.success("Document successfully generated and uploaded to HackMD!")
            
            st.markdown("#### HackMD Link:")
            link_col1, link_col2 = st.columns([3, 1])
            
            with link_col1:
                st.text_area("Copy the link below:", hackmd_link, height=70, key="hackmd_link_text")
            
            with link_col2:
                st.markdown(f"[Open in Browser]({hackmd_link})")
            
            # If this was a split document, provide additional information
            if "Part" in hackmd_link:
                st.info("""
                Note: Due to the large size of your document, it was automatically split into multiple parts.
                The link above points to an index page that contains links to all parts of your document.
                """)
                
        return hackmd_link


class AssignmentDocumentGenerator(DocumentGenerator):
    """Generator for Assignment Reference Documents"""
    
    def generate(self, content, is_direct_input=False) -> Optional[str]:
        """Generate Assignment Reference Document from file or direct input"""
        if not self._validate_parameters():
            return None
            
        document_type = "Assignment Reference Document"
        
        # Progress bar
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        try:
            # Extract text
            status_text.text('Reading and analyzing content...')
            progress_bar.progress(10)
            
            if is_direct_input:
                # Use the direct input text
                pdf_text = content
            else:
                # Extract from uploaded file
                file_type = content.name.split(".")[-1].lower()
                pdf_text = TextExtractor.extract(content, file_type)
                
            if not pdf_text or pdf_text.strip() == "":
                st.error("Failed to extract text from the content.")
                return None
                
            progress_bar.progress(30)
            status_text.text('Text extracted successfully!')
            
            # Analyze text
            status_text.text('Preparing Reference Document...')
            prompt = f""" 
     You are tasked with creating a detailed reference document based on the following assignment description:

    Assignment Description:
    {pdf_text}

    Create a document with the following formatting and structure:

    1. Use main section headings as bold side headings in the format "**Heading:**"
    2. Start directly with the Objective section (no title, no introduction)
    3. Each major section should have its heading in bold followed by a colon

    **Objective:** Provide a concise overview of the assignment's purpose and goals. Clearly explain the main objectives and expected learning outcomes.

    **Step-by-Step Instructions:** Create detailed, easy-to-follow steps for completing the assignment. Each step should be actionable and include:
    
    **Project Setup and Initialization:** Describe how to set up the project directory, initialize the project, and install dependencies.
    
    **Development Process:** Outline detailed steps for Briefly Explained in detail way in lengthy How to develop the Assignment.
    
    **Styling and Design:** Provide instructions for implementing styling.
    
    **Deployment:** Describe the process of deploying the project to a suitable platform and configuring the environment.

    **Submission Guidelines:** Clearly define how the assignment should be submitted, including the expected format, any required documentation.

    *Important Notes:*
    - Format all major section headings as "**Heading:**" (bold with colon)
    - Make subsection headings bold with colon as well
    - Do not include a title like "Reference Document" or "Project Guide"
    - Start directly with "**Objective:**" as the first line of the document
    - Do not include any code snippets
    - Do not include a summary or conclusion

    Generate the content in a structured and professional manner, ensuring clarity and usability.
    """
            # Create rate limit handler for API call
            rate_limiter = RateLimitHandler(
                requests_per_minute=15,  # Gemini's limit
                retry_delay=65,          # Wait a bit over a minute to be safe
                max_retries=5            # Allow up to 5 retries
            )
            
            # Use rate limiter for API call
            analyzed_text = rate_limiter.execute_with_rate_limit(analyze_text_with_gemini, prompt)
            
            if not analyzed_text or analyzed_text.strip() == "":
                st.error("Failed to analyze text. Please try again.")
                return None
                
            progress_bar.progress(70)
            status_text.text('Document content generated successfully!')
            
            # Upload and log
            status_text.text('Uploading to HackMD...')
            hackmd_link = self.upload_and_log(analyzed_text, document_type)
            
            progress_bar.progress(100)
            status_text.text('Process completed successfully!')
            
            return analyzed_text
            
        except Exception as e:
            logger.error(f"Error generating assignment document: {str(e)}")
            st.error(f"Error generating document: {str(e)}")
            return None


class ConceptsDocumentGenerator(DocumentGenerator):
    """Generator for Concepts Reference Documents with rate limit handling"""
    
    def generate(self, content, is_direct_input=False) -> Optional[str]:
        """Generate Concepts Reference Document with rate limit handling"""
        if not self._validate_parameters():
            return None
            
        document_type = "Concepts Reference Document"
        
        # Progress bar
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        try:
            # Extract text from all files or use direct input
            status_text.text('Reading and analyzing content...')
            
            if is_direct_input:
                # Use direct text input
                extracted_text = content
                progress_bar.progress(30)
            else:
                # Extract from uploaded files
                extracted_text = ""
                
                # Process files in parallel
                with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
                    futures = []
                    for uploaded_file in content:
                        file_type = uploaded_file.name.split(".")[-1].lower()
                        futures.append(
                            executor.submit(TextExtractor.extract, uploaded_file, file_type)
                        )
                    
                    # Collect results as they complete
                    total_files = len(futures)
                    for i, future in enumerate(concurrent.futures.as_completed(futures)):
                        try:
                            file_text = future.result()
                            if file_text and file_text.strip():
                                extracted_text += file_text + "\n"
                            progress_value = min(30, int(30 * (i + 1) / total_files))
                            progress_bar.progress(progress_value)
                        except Exception as e:
                            logger.error(f"Error processing file {i}: {str(e)}")
                            st.warning(f"Error processing one of the files: {str(e)}")
            
            if not extracted_text.strip():
                st.error("Failed to extract any text. Please check the content.")
                return None
                
            status_text.text('Content read successfully!')
            
            # Create rate limit handler
            rate_limiter = RateLimitHandler(
                requests_per_minute=15,  # Gemini's limit
                retry_delay=65,          # Wait a bit over a minute to be safe
                max_retries=10           # Allow up to 10 retries
            )
            
            # Analyze concepts with rate limiting
            status_text.text('Analyzing concepts...')
            prompt = f"""
Please follow the steps below for analyzing the extracted content:

1. **Analyze Questions from Images**:
- First, focus on identifying and extracting any questions from the images or text. A question is defined as a sentence or phrase that is asking for specific information, typically ending with a question mark.
- If the extracted text includes any question-like phrases, please isolate them and remove any surrounding context.
- Remove unnecessary prefixes like "Asked about," "Can you explain," "What is," or similar phrases. 
- **If there is any question or concept related to 'Introduce Yourself,' remove it entirely from the analysis.**
- Ignore any broad or generic questions such as:
    - Questions about educational background (e.g., "What was your major technology focus during BE?", "Technologies Used In Current Projects")
    - Questions about personal details or self-introduction (e.g., "Can you introduce yourself?")
    - Resume-related questions (e.g., "Rate your skills as mentioned on your resume")
    - Generic or vague questions that do not focus on specific technical concepts

2. **Extract Key Concepts**:
- Convert questions into key concepts or topics they are related to.
- Each concept should be a descriptive phrase that directly reflects the essence of the question.
- Use longer, descriptive names for concepts to ensure clarity and context. For example:
    - Instead of "Polymorphism," use "Polymorphism in Object-Oriented Programming."
    - Instead of "HTTP," use "HTTP Protocol and Its Methods."

3. **Filter Unique Concepts**:
- Remove duplicate concept names.
- Ensure there are no repeated entries in the list.

4. **Add Similar Concepts**:
- Based on the extracted unique concepts, generate 10 additional relevant concepts aligned with the context.

5. **Display Format - VERY IMPORTANT**:
- Return ONLY a simple list of concept names, one per line
- DO NOT include ANY explanatory text, headers, or context
- DO NOT include ANY numbering or bullet points
- DO NOT prefix concepts with "Concept:" or similar text
- Just provide clean concept names, one per line, each starting with a capital letter
- NO explanations, NO prefixes, NO context, ONLY the concept names themselves

Here is the content to analyze:

{extracted_text}
"""
            
            # Use rate limiter for API call
            analyzed_concepts = rate_limiter.execute_with_rate_limit(analyze_text_with_gemini, prompt)
            
            if not analyzed_concepts or analyzed_concepts.strip() == "":
                st.error("Failed to analyze concepts. Please verify the content.")
                return None
                
            progress_bar.progress(50)
            status_text.text('Concepts analyzed successfully!')
            
            # Display concepts preview
            concepts_list = analyzed_concepts.split("\n")
            filtered_concepts = [
                concept.strip() 
                for concept in concepts_list 
                if concept.strip() and "?" not in concept
            ]
            
            # Show the number of concepts found
            num_concepts = len(filtered_concepts)
            if num_concepts > 15:
                st.warning(f"Found {num_concepts} concepts. This may exceed Gemini's rate limit of 15 requests per minute. The process will automatically pause when needed and resume after the rate limit resets.")
            
            st.markdown("### Preview of Extracted Concepts:")
            preview_container = st.container()
            with preview_container:
                for concept in filtered_concepts:
                    st.write(f"- {concept}")
            
            # Option to select only specific concepts for processing
            if len(filtered_concepts) > 5:
                st.info("If you'd like to process only specific concepts, you can select them below. Otherwise, all concepts will be processed.")
                selected_concepts = st.multiselect(
                    "Select specific concepts to process (optional)",
                    filtered_concepts,
                    default=None
                )
                
                if selected_concepts:
                    filtered_concepts = selected_concepts
                    st.success(f"Processing {len(filtered_concepts)} selected concepts.")
            
            # Generate detailed content for each concept
            status_text.text('Generating detailed explanations for concepts...')
            full_document_content = ""
            
            # Store already processed concepts in case of interruption
            processed_concepts = {}
            concept_status = st.empty()
            
            # Create placeholder for rate limit info
            rate_limit_info = st.empty()
            
            # Process concepts sequentially with rate limiting
            for i, concept in enumerate(filtered_concepts):
                concept_number = i + 1
                concept_status.info(f"Processing concept {concept_number}/{len(filtered_concepts)}: {concept}")
                
                solution_prompt = f"""
                Provide a theoretical answer for the concept: {concept}, ensuring key details are covered without including a summary or conclusion.

                **For Coding-Related Questions**:
                - Do NOT provide complete code.
                - Instead, describe the approach in a structured manner:
                - Explain the problem briefly.
                - Outline the key steps needed to solve it.
                - Mention important considerations like time complexity, edge cases, and best practices.

                **For Non-Coding Theoretical Concepts**:
                - Provide an in-depth explanation covering relevant details.
                - Use bullet points or structured formatting where appropriate.

                Concept: {concept}
                """
                
                try:
                    # Use rate limiter for API call
                    solution = rate_limiter.execute_with_rate_limit(analyze_text_with_gemini, solution_prompt)
                    
                    if solution:
                        processed_concepts[concept] = solution
                        full_document_content += f"### {concept}\n\n{solution}\n\n"
                        
                    # Update progress
                    progress_value = min(80, 50 + int(30 * (i + 1) / len(filtered_concepts)))
                    progress_bar.progress(progress_value)
                    status_text.text(f'Generated {i+1}/{len(filtered_concepts)} concepts...')
                    
                    # Show remaining rate limit info
                    remaining_requests = 15 - len(rate_limiter.request_timestamps)
                    rate_limit_info.info(f"API calls remaining before rate limit: {max(0, remaining_requests)}/15")
                    
                except Exception as e:
                    logger.error(f"Error generating solution for concept '{concept}': {str(e)}")
                    st.warning(f"Error generating solution for concept '{concept}': {str(e)}")
                    # Continue with next concept rather than stopping completely
                    continue
            
            # Upload and log
            if full_document_content:
                status_text.text('Uploading to HackMD...')
                hackmd_link = self.upload_and_log(full_document_content, document_type)
                
                progress_bar.progress(100)
                status_text.text('Process completed successfully!')
                
                # Clear temporary status indicators
                concept_status.empty()
                rate_limit_info.empty()
                
                return full_document_content
            else:
                st.error("Failed to generate any content for concepts. Please try again.")
                return None
                
        except Exception as e:
            logger.error(f"Error generating concepts document: {str(e)}")
            st.error(f"Error generating document: {str(e)}")
            return None


class BatchConceptsDocumentGenerator(DocumentGenerator):
    """Generator for Concepts Reference Documents using batch processing"""
    
    def generate(self, content, is_direct_input=False) -> Optional[str]:
        """Generate Concepts Reference Document with batch processing"""
        if not self._validate_parameters():
            return None
            
        document_type = "Concepts Reference Document"
        
        # Progress bar
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        try:
            # Extract text from all files or use direct input
            status_text.text('Reading and analyzing content...')
            
            if is_direct_input:
                # Use direct text input
                extracted_text = content
                progress_bar.progress(30)
            else:
                # Extract from uploaded files
                extracted_text = ""
                
                with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
                    futures = []
                    for uploaded_file in content:
                        file_type = uploaded_file.name.split(".")[-1].lower()
                        futures.append(
                            executor.submit(TextExtractor.extract, uploaded_file, file_type)
                        )
                    
                    total_files = len(futures)
                    for i, future in enumerate(concurrent.futures.as_completed(futures)):
                        try:
                            file_text = future.result()
                            if file_text and file_text.strip():
                                extracted_text += file_text + "\n"
                            progress_value = min(30, int(30 * (i + 1) / total_files))
                            progress_bar.progress(progress_value)
                        except Exception as e:
                            logger.error(f"Error processing file {i}: {str(e)}")
                            st.warning(f"Error processing one of the files: {str(e)}")
            
            if not extracted_text.strip():
                st.error("Failed to extract any text. Please check the content.")
                return None
                
            status_text.text('Content read successfully!')
            
            # Analyze concepts
            status_text.text('Analyzing concepts...')
            prompt = f"""
            Please follow the steps below for analyzing the extracted content:

            1. **Analyze Questions from Images**:
            - First, focus on identifying and extracting any questions from the images. A question is defined as a sentence or phrase that is asking for specific information, typically ending with a question mark.
            - If the extracted text includes any question-like phrases from images, please isolate them and remove any surrounding context.
            - Remove unnecessary prefixes like "Asked about," "Can you explain," "What is," or similar phrases. 
            - **If there is any question or concept related to 'Introduce Yourself,' remove it entirely from the analysis.**
            - Ignore any broad or generic questions such as:
                - Questions about educational background (e.g., "What was your major technology focus during BE?", "Technologies Used In Current Projects")
                - Questions about personal details or self-introduction (e.g., "Can you introduce yourself?")
                - Resume-related questions (e.g., "Rate your skills as mentioned on your resume")
                - Generic or vague questions that do not focus on specific

            2. **Extract Key Concepts**:
            - Once the questions have been analyzed and extracted, identify the key concepts or topics they are related to. Each concept should be a descriptive phrase that directly reflects the essence of the question. Avoid using overly generic terms.
            - Use longer, descriptive names for concepts to ensure clarity and context. For example:
                - Instead of "Polymorphism," use "Polymorphism in Object-Oriented Programming."
                - Instead of "HTTP," use "HTTP Protocol and Its Methods."

            3. **Filter Unique Concepts**:
            - Remove duplicate concept names. Display only unique, descriptive concepts, ensuring there are no repeated entries in the list.

            4. *Add Similar Concepts*:
            - Based on the extracted unique concepts, generate 10 additional descriptive and relevant concepts aligned with the context.

            5. **Provide Preview**:
            - Display the list of unique extracted concepts followed by the list of similar concepts. These should be displayed as a simple list with each concept on a new line and each concept name starting with a capital letter.
            - **Do not include any extra context, such as headers like "Key Concepts" or "Extracted Questions." Just show the concept names cleanly.**

            Here is the content to analyze:

            {extracted_text}
            """
            
            # We'll still make a single call for the concept extraction
            analyzed_concepts = analyze_text_with_gemini(prompt)
            
            if not analyzed_concepts or analyzed_concepts.strip() == "":
                st.error("Failed to analyze concepts. Please verify the content.")
                return None
                
            progress_bar.progress(40)
            status_text.text('Concepts analyzed successfully!')
            
            # Display concepts preview
            concepts_list = analyzed_concepts.split("\n")
            filtered_concepts = [
                concept.strip() 
                for concept in concepts_list 
                if concept.strip() and "?" not in concept
            ]
            
            # Show the number of concepts found
            num_concepts = len(filtered_concepts)
            if num_concepts > 15:
                st.warning(f"Found {num_concepts} concepts. Using batch processing to handle them efficiently.")
            
            st.markdown("### Preview of Extracted Concepts:")
            for concept in filtered_concepts:
                st.write(f"- {concept}")
                
            # Option to select only specific concepts for processing
            if len(filtered_concepts) > 5:
                st.info("You can select specific concepts to process or process all of them.")
                
                # Option to process all concepts or select specific ones
                process_mode = st.radio(
                    "Select processing mode:",
                    ["Process all concepts", "Select specific concepts"]
                )
                
                if process_mode == "Select specific concepts":
                    selected_concepts = st.multiselect(
                        "Select concepts to process",
                        filtered_concepts,
                        default=filtered_concepts[:5]  # Default to first 5
                    )
                    
                    if selected_concepts:
                        filtered_concepts = selected_concepts
                        st.success(f"Processing {len(filtered_concepts)} selected concepts.")
            
            # Ask for batch size
            col1, col2 = st.columns(2)
            with col1:
                batch_size = st.number_input(
                    "Concepts per batch (recommended: 10-15 for Gemini free tier)",
                    min_value=1,
                    max_value=50,
                    value=10
                )
            
            with col2:
                batch_delay = st.number_input(
                    "Seconds to wait between batches",
                    min_value=30,
                    max_value=120,
                    value=65
                )
            
            # Initialize batch processor
            def update_progress(current, total):
                progress_value = min(80, 40 + int(40 * current / total))
                progress_bar.progress(progress_value)
                status_text.text(f'Generated {current}/{total} concepts...')
            
            batch_processor = BatchConceptsProcessor(
                batch_size=batch_size,
                batch_delay=batch_delay,
                api_function=analyze_text_with_gemini,
                progress_callback=update_progress
            )
            
            if st.button("Start Processing Concepts"):
                # Generate detailed content for each concept using batch processing
                status_text.text('Generating detailed explanations for concepts...')
                
                # Process concepts in batches
                concept_results = batch_processor.process_concepts(
                    concepts=filtered_concepts,
                    prompt_template=solution_prompt_for_concept
                )
                
                # Generate the full document content
                full_document_content = ""
                
                for concept in filtered_concepts:
                    if concept in concept_results:
                        solution = concept_results[concept]
                        full_document_content += f"### {concept}\n\n{solution}\n\n"
                
                # Upload and log
                if full_document_content:
                    status_text.text('Uploading to HackMD...')
                    hackmd_link = self.upload_and_log(full_document_content, document_type)
                    
                    progress_bar.progress(100)
                    status_text.text('Process completed successfully!')
                    
                    return full_document_content
                else:
                    st.error("Failed to generate any content for concepts. Please try again.")
                    return None
            
            return None
                
        except Exception as e:
            logger.error(f"Error generating concepts document: {str(e)}")
            st.error(f"Error generating document: {str(e)}")
            return None


def generate_reference_document_to_docx(analyzed_text, output_filename_docx, output_stream):
    """Generate a formatted DOCX document from markdown text"""
    doc = Document()

    paragraphs = analyzed_text.split("\n\n")  

    for paragraph in paragraphs:
        lines = paragraph.splitlines()
        for line in lines:
            if line.startswith('#### '):
                p = doc.add_paragraph(line[5:])
                p.style = 'Heading 4'
            elif line.startswith('### '):
                p = doc.add_paragraph(line[4:])
                p.style = 'Heading 3'
            elif line.startswith('## '):
                p = doc.add_paragraph(line[3:])
                p.style = 'Heading 2'
            elif line.startswith('# '):
                p = doc.add_paragraph(line[2:])
                p.style = 'Heading 1'
            else:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])  
                        run.bold = True
                    else:
                        p.add_run(part)

    doc.save(output_stream)


def show_usage_instructions():
    """Show instructions for using the app"""
    with st.expander("How to use this app"):
        st.markdown("""
        ### Assignment Reference Document Generator
        
        This tool helps you generate reference documents for assignments:
        
        1. **Select document type**: Choose "Assignment Reference Document"
        2. **Fill in the details**: Title, Job ID, your name, and job type
        3. **Choose input method**: Upload files or directly input text
        4. **Generate**: Click the "Generate Document" button
        5. **Get link**: The tool will generate a document and provide a HackMD link
        
        ### Concepts Reference Document Generator
        
        This tool analyzes interview questions and generates explanations:
        
        1. **Select document type**: Choose "Concepts Reference Document"
        2. **Choose processing mode**: Standard or Batch processing
        3. **Fill in the details**: Title, Job ID, your name, job type, and interview round
        4. **Choose input method**: Upload files or directly input text
        5. **Generate**: Click the "Generate Document" button
        6. **Review concepts**: The tool will extract concepts for your review
        7. **Get link**: A HackMD link with all explanations will be provided
        """)


# The key issue is in the reference_document_generator function
# We need to move the input method selection outside the form
# and use session state to maintain the selection

def reference_document_generator():
    """Main Streamlit UI function for the Reference Document Generator"""
    st.header("Reference Document Generator")
    
    # Add styling
    st.markdown("""
    <style>
    .stButton button {
        background-color: #4CAF50;
        color: white;
        padding: 10px 20px;
        border: none;
        border-radius: 4px;
        cursor: pointer;
    }
    .stButton button:hover {
        background-color: #45a049;
    }
    .warning {
        color: red;
        font-weight: bold;
    }
    </style>
    """, unsafe_allow_html=True)

    # App tabs
    tab1, tab2 = st.tabs(["Generate Document", "Download Previous Document"])
    
    with tab1:
        # Document type selection
        document_type = st.selectbox("Select Document Type", ["Assignment Reference Document", "Concepts Reference Document"])

        # For Concepts Reference Document, add processing mode selection
        processing_mode = None
        if document_type == "Concepts Reference Document":
            processing_mode = st.radio(
                "Select processing mode:",
                ["Standard (Automatic pause and retry)", "Batch processing (Manual control)"]
            )

        # Move input method selection outside the form
        # Initialize session state for input method if it doesn't exist
        if 'input_method' not in st.session_state:
            st.session_state.input_method = "Upload Files"
            
        # Create the input method selector outside the form
        input_method = st.radio(
            "Select Input Method",
            ["Upload Files", "Direct Text Input"],
            help="Choose whether to upload files or input text directly",
            key="input_method"  # This links it to session state
        )

        # Input fields with validation
        with st.form(key="document_form"):
            title_name = st.text_input("Enter Title Name for the Document", 
                                    placeholder="Enter the title to be displayed in HackMD",
                                    help="This will be displayed as the title in HackMD")
            
            job_id = st.text_input("Enter Job ID", 
                                placeholder="Enter the Job ID associated with this document",
                                help="This will be used for tracking in the Google Sheet")
            
            prepared_by = st.text_input("Prepared By", 
                                    placeholder="Enter your name",
                                    help="Your name will be recorded in the logs")
            
            job_type = st.selectbox("Select Job Type", 
                                ["Full Time", "Internships", "Intern+FullTime"],
                                help="Type of job this document is for")

            # Show Interview Round dropdown for "Concepts Reference Document"
            interview_round = None
            if document_type == "Concepts Reference Document":
                interview_round = st.selectbox(
                    "Select Interview Round",
                    [
                        "Assessment", "TR1", "TR2", "MR", "HR", 
                        "TR1 + TR2 + HR", "Assessment + TR1 + TR2 + HR", 
                        "Offline drive", "CEO Round", "Culture Fit Round"
                    ],
                    help="The interview round this document is for"
                )

            # File Uploader or Text Input based on selected method
            if st.session_state.input_method == "Upload Files":
                if document_type == "Assignment Reference Document":
                    st.write("### Upload Files")
                    st.write("Please upload the assignment document (PDF, DOCX, or text file)")
                    uploaded_files = st.file_uploader(
                        "Upload assignment document", 
                        type=["pdf", "docx", "txt"], 
                        accept_multiple_files=False,
                        help="Upload a single file containing the assignment description"
                    )
                    direct_input_text = None
                else:  # Concepts Reference Document
                    st.write("### Upload Files")
                    st.write("Please upload interview questions or concepts (PDF, DOCX, ZIP, Image files)")
                    st.write("For best results, ensure images are clear and readable")
                    uploaded_files = st.file_uploader(
                        "Upload files", 
                        type=SUPPORTED_FILE_TYPES, 
                        accept_multiple_files=True,
                        help="Upload one or more files containing interview questions or concepts"
                    )
                    direct_input_text = None
            else:  # Direct Text Input
                st.write("### Enter Text Directly")
                if document_type == "Assignment Reference Document":
                    direct_input_text = st.text_area(
                        "Enter assignment description",
                        height=300,
                        placeholder="Paste or type the assignment description here...",
                        help="Type or paste the assignment description directly"
                    )
                else:  # Concepts Reference Document
                    direct_input_text = st.text_area(
                        "Enter interview questions or concepts",
                        height=300,
                        placeholder="Paste or type the interview questions or concepts here...",
                        help="Type or paste the interview questions or concepts directly"
                    )
                uploaded_files = None
            
            submit_button = st.form_submit_button("Generate Document")
        
        # Process form submission
        if submit_button:
            # Basic validation
            if st.session_state.input_method == "Upload Files" and not uploaded_files:
                st.error("Please upload at least one file before proceeding.")
                st.stop()
            elif st.session_state.input_method == "Direct Text Input" and not direct_input_text:
                st.error("Please enter some text before proceeding.")
                st.stop()
                
            # For Assignment Reference Document
            if document_type == "Assignment Reference Document":
                generator = AssignmentDocumentGenerator(
                    title_name, 
                    job_id,
                    prepared_by,
                    job_type
                )
                
                if st.session_state.input_method == "Upload Files":
                    generator.generate(uploaded_files, is_direct_input=False)
                else:  # Direct Text Input
                    generator.generate(direct_input_text, is_direct_input=True)
                
            # For Concepts Reference Document
            elif document_type == "Concepts Reference Document":
                if st.session_state.input_method == "Upload Files" and not isinstance(uploaded_files, list):
                    uploaded_files = [uploaded_files]
                
                # Choose generator based on processing mode
                if processing_mode == "Standard (Automatic pause and retry)":
                    generator = ConceptsDocumentGenerator(
                        title_name,
                        job_id,
                        prepared_by,
                        job_type,
                        interview_round
                    )
                else:  # Batch processing
                    generator = BatchConceptsDocumentGenerator(
                        title_name,
                        job_id,
                        prepared_by,
                        job_type,
                        interview_round
                    )
                
                if st.session_state.input_method == "Upload Files":
                    generator.generate(uploaded_files, is_direct_input=False)
                else:  # Direct Text Input
                    generator.generate(direct_input_text, is_direct_input=True)
    
    with tab2:
        st.write("### Download Previously Generated Document")
        st.write("Enter the HackMD link of the document you want to download:")
        
        hackmd_url = st.text_input("HackMD URL", placeholder="https://hackmd.io/...")
        
        if st.button("Download as DOCX") and hackmd_url:
            with st.spinner("Fetching document and converting to DOCX..."):
                try:
                    # Extract document ID from URL
                    doc_id = hackmd_url.split("/")[-1]
                    
                    # Fetch content from HackMD API
                    api_url = f"https://api.hackmd.io/v1/notes/{doc_id}"
                    headers = {"Authorization": f"Bearer {HACKMD_API_TOKEN}"}
                    
                    response = requests.get(api_url, headers=headers)
                    if response.status_code == 200:
                        content = response.json().get("content", "")
                        title = response.json().get("title", "Document")
                        
                        # Generate DOCX
                        output = BytesIO()
                        generate_reference_document_to_docx(content, f"{title}.docx", output)
                        output.seek(0)
                        
                        # Offer download
                        st.download_button(
                            label="Download DOCX",
                            data=output,
                            file_name=f"{title}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        st.error(f"Failed to fetch document. Status code: {response.status_code}")
                except Exception as e:
                    logger.error(f"Error downloading document: {str(e)}")
                    st.error(f"Error downloading document: {str(e)}")


def main():
    """Main function to run the Streamlit app"""
    # Configure page settings
    st.set_page_config(
        page_title="Reference Document Generator",
        page_icon="📄",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # App header with logo and title
    col1, col2 = st.columns([1, 5])
    with col1:
        st.image("https://res.cloudinary.com/dg8n2jeur/image/upload/v1741718114/cws2futayeddahz3sz8z.webp", width=80)
    with col2:
        st.title("Reference Document Generator")
        st.markdown("Create professional reference documents for assignments and interview concepts")
    
    # Sidebar information
    with st.sidebar:
        st.header("About")
        st.info(
            "This application helps create structured reference documents "
            "for assignments and technical interview concepts. Documents are "
            "automatically uploaded to HackMD and logged in Google Sheets."
        )
        
        st.header("Options")
        # Add theme selector or other preferences here
        theme = st.selectbox("Theme", ["Light", "Dark"])
        if theme == "Dark":
            # Apply dark theme CSS
            st.markdown("""
            <style>
            body {
                background-color: #121212;
                color: #ffffff;
            }
            </style>
            """, unsafe_allow_html=True)
        
        # Show rate limit information
        st.header("API Rate Limit Info")
        st.warning(
            "Gemini API has a rate limit of 15 requests per minute. "
            "When generating solutions for many concepts, the application "
            "will automatically pause when needed and resume after the rate limit resets."
        )
        
        st.info(
            "For large numbers of concepts, consider using the 'Batch processing' mode "
            "which allows you to configure batch sizes and processing intervals."
        )
        

    # Check dependencies
    dependency_error = False
    try:
        import gspread
        import requests
        import pytesseract
        import google.generativeai
    except ImportError as e:
        st.error(f"Missing dependency: {str(e)}")
        dependency_error = True
    
    if dependency_error:
        st.warning("Some features may not work properly due to missing dependencies.")
    
    # Show instruction expander
    show_usage_instructions()
    
    # Main app
    reference_document_generator()
    
    # Footer
    st.markdown("---")
    st.markdown(
        "© 2025 Reference Document Generator | "
        "For support, contact tirupathirao.kella@nxtwave.co.in"
    )


if __name__ == "__main__":
    # Set up exception handling
    try:
        main()
    except Exception as e:
        st.error(f"An unexpected error occurred: {str(e)}")
        logger.critical(f"Critical application error: {str(e)}", exc_info=True)
